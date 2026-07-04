# -*- coding: utf-8 -*-
"""
生成 Neuroelectrics Enobio & Starstim 引进中国国产化注册可行性评估报告
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_run_font(run, name='微软雅黑', size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rpr = r.rPr
    if rpr is None:
        rpr = OxmlElement('w:rPr')
        r.insert(0, rpr)
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def add_heading(doc, text, level=1):
    styles = {
        1: (16, True, (31, 73, 125)),
        2: (13.5, True, (54, 95, 145)),
        3: (12, True, (68, 114, 148)),
    }
    size, bold, color = styles.get(level, (11, True, (0, 0, 0)))
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_para(doc, text, bold=False, indent_cn=True, size=10.5, color=None, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent_cn:
        p.paragraph_format.first_line_indent = Pt(21)
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=10.5, bold=True)
        r2 = p.add_run(text)
        set_run_font(r2, size=10.5)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.5)
    return p


def add_table(doc, headers, rows, col_widths=None, header_color='1F497D'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        set_cell_bg(hdr[i], header_color)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ''
            p = cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            set_run_font(run, size=9.5)
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


# ============ 开始生成文档 ============
doc = Document()

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# 默认字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ==== 封面 ====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(120)
r = p.add_run('Neuroelectrics 脑电与经颅电刺激产品')
set_run_font(r, size=22, bold=True, color=(31, 73, 125))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('引进中国"国产化"注册可行性深度评估报告')
set_run_font(r, size=20, bold=True, color=(31, 73, 125))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
r = p.add_run('—— 以 Enobio（EEG）与 Starstim（tES+EEG）为分析对象 ——')
set_run_font(r, size=13, color=(89, 89, 89))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(140)
r = p.add_run('注册路径 · 时间成本 · 费用测算 · 对方报价合理性判定')
set_run_font(r, size=12, color=(89, 89, 89))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
r = p.add_run('评估周期：2026 年 7 月')
set_run_font(r, size=11, color=(89, 89, 89))

doc.add_page_break()

# ==== 目录（手动生成，Word 打开后可 F9 刷新真实目录） ====
add_heading(doc, '目  录', level=1)
toc_items = [
    '一、执行摘要（Executive Summary）',
    '二、产品与厂家背景',
    '    2.1 Neuroelectrics 公司概况',
    '    2.2 Enobio 产品线（EEG）',
    '    2.3 Starstim 产品线（tES + EEG）',
    '三、中国 NMPA 法规环境与产品分类判定',
    '    3.1 医疗器械分类界定',
    '    3.2 进口与"国产"两种路径的本质差异',
    '    3.3 免于临床试验目录（豁免临床）适用性分析',
    '四、"以国产形式申报"的落地方案',
    '    4.1 三种国产化模式对比',
    '    4.2 生产场地与体系建设要求',
    '    4.3 关键前提：与西班牙方的技术转让/授权协议',
    '五、注册路径与临床策略详解',
    '    5.1 Enobio（EEG）：对标同品种、豁免临床路径',
    '    5.2 Starstim（tES）：必须做临床的路径',
    '    5.3 同步 vs 依次策略的定量对比',
    '六、时间进度评估',
    '    6.1 关键里程碑与关键路径',
    '    6.2 三种策略时间对比表',
    '七、费用测算与拆解',
    '    7.1 Enobio 单独注册费用（无临床）',
    '    7.2 Starstim 带临床注册全成本',
    '    7.3 与西班牙代表报价的逐项比对',
    '八、对方报价合理性判定',
    '    8.1 100–150 万"无临床注册费"合理性',
    '    8.2 150 万设备成本、23 万耗材合理性',
    '    8.3 250 万医生费用合理性（重点关注）',
    '    8.4 同步 +150 万 / 依次省设备的策略合理性',
    '九、风险清单与不确定性提示',
    '十、结论与行动建议',
    '附录 A：假设与数据来源说明',
    '附录 B：术语表',
]
for it in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(it)
    set_run_font(r, size=11)

doc.add_page_break()

# ==== 一、执行摘要 ====
add_heading(doc, '一、执行摘要（Executive Summary）', level=1)

add_para(doc,
    '本报告基于 Neuroelectrics 公司提供的 Enobio（无线干电极脑电采集系统）与 Starstim'
    '（脑电采集 + 经颅电刺激一体机）两款产品资料，结合中国 NMPA（国家药品监督管理局）'
    '现行医疗器械注册法规，对以下核心问题进行了系统性评估：'
    '（1）"以中国国产形式申报"的可行路径；'
    '（2）"对标同品种以豁免/减免临床"路径的适用性；'
    '（3）时间、费用、流程的合理区间；'
    '（4）西班牙方中国代表所报出的费用组合是否合理。'
)

add_heading(doc, '核心结论（结论前置）', level=2)

concl_rows = [
    ['结论 1', '产品分类', 'Enobio 定为二类；Starstim 因含经颅电刺激功能，若声称治疗用途按三类管理，若仅用于科研/BCI 声称则可能按二类；本报告按"三类治疗用途"假设推进（对方要走临床即隐含此假设）。'],
    ['结论 2', 'Enobio 走豁免临床合理', '脑电图机在《免于临床评价目录》内，走同品种对比路径 100–180 万人民币、12–18 个月可拿证。对方 100–150 万报价基本落在合理区间下限，属"熟练代理公司报价"。'],
    ['结论 3', 'Starstim 必须做临床', '经颅电刺激类不在豁免目录内，必须做多中心 RCT。总花费保守估计 550–950 万人民币（含设备/耗材/临床/注册），周期 30–48 个月。对方给出的 150 万设备 + 23 万耗材 + 250 万医生的组合价格偏低。'],
    ['结论 4', '250 万医生费用偏低', '按 3 中心、100–150 例、随访 3–6 个月的常规三类神经调控临床，PI/CRC/受试者补偿/伦理/统计/CRO 合计 400–800 万为常见区间。250 万只能覆盖"研究者费用"或走"关系型压价 + 论文归属激励"通道。存在可行性但需绑定强关系资源，且论文归属激励不能替代 GCP 合规成本。'],
    ['结论 5', '同步 vs 依次策略', '"同步做 +150 万设备成本"是合理的表述（多一套 15–20 台设备）；但同步策略的真实增量应≈180–220 万（含额外 CRC/监查/耗材），单说 150 万偏乐观。依次策略节省 150 万但延迟 12–18 个月上市，机会成本通常高于节省额。'],
    ['结论 6', '"国产化"是关键前提', '想真正以"国产"身份申报，必须在中国境内完成实质性生产（至少关键工序），并取得境外原厂的技术授权/OEM 协议。这笔"隐性成本"（体系建立、场地、授权费）对方报价中并未体现，需单独测算 300–800 万人民币或走中国 OEM 代工路径。'],
]
add_table(doc, ['编号', '要点', '说明'], concl_rows, col_widths=[1.5, 3.5, 11])

add_para(doc,
    '综上：对方代表的报价在"注册代理"这一层是合理的，但没有覆盖"国产化落地"与"临床合规全成本"'
    '两块隐性支出。建议以本报告的费用与时间区间作为谈判和内部预算的基准，并在合同层面锁定'
    '"技术转让 + 生产授权 + 临床数据可用性"三项关键条款。'
)

doc.add_page_break()

# ==== 二、产品与厂家背景 ====
add_heading(doc, '二、产品与厂家背景', level=1)

add_heading(doc, '2.1 Neuroelectrics 公司概况', level=2)
add_para(doc,
    'Neuroelectrics（总部：西班牙巴塞罗那）成立于 2011 年，是全球少数同时拥有"高密度无线脑电"'
    '（EEG）与"多通道经颅电刺激"（tES：tDCS/tACS/tRNS）两条产品线的公司。其代表产品：'
)
add_bullet(doc, 'Enobio 系列：面向科研与临床的无线 EEG 采集系统，8/20/32 通道，最新推出 EnobioDx 面向临床试验场景（2026 版资料）。')
add_bullet(doc, 'Starstim 系列：兼具 EEG 采集与 tES 刺激功能，代表机型 Starstim 8/32/tES，可实现闭环神经调控（closed-loop neuromodulation），面向卒中康复、抑郁、认知障碍等临床研究与治疗探索。')
add_para(doc,
    '两条产品线均通过欧盟 CE 认证，Starstim 部分型号在美国有 FDA 510(k)（用于研究用途 IDE）。'
    '在中国目前无 NMPA 注册证，属首次进入中国市场。'
)

add_heading(doc, '2.2 Enobio 产品线（EEG）', level=2)
enobio_rows = [
    ['产品定位', '无线干/湿电极脑电采集，8/20/32 通道'],
    ['预期用途（拟报）', '临床脑电信号采集与显示，辅助神经科诊断/研究'],
    ['风险等级参照', '中国《医疗器械分类目录》06-09 神经和肌肉刺激/06-10 电生理，脑电图机为 二类（Class II）'],
    ['同类已上市产品', '国内已有多家（迈瑞、上海诺诚、中科新松、Compumedics 等）取得二类脑电图机注册证'],
    ['临床评价路径', '列入《免于临床评价医疗器械目录》，可走同品种对比路径豁免临床'],
]
add_table(doc, ['项目', '内容'], enobio_rows, col_widths=[4, 12])

add_heading(doc, '2.3 Starstim 产品线（tES + EEG）', level=2)
starstim_rows = [
    ['产品定位', '经颅电刺激（tDCS/tACS/tRNS）+ 同步 EEG 采集，8/20/32 通道'],
    ['预期用途（关键分歧点）', '（A）若拟报"抑郁/卒中康复/疼痛等适应症治疗" → 按 三类（Class III）管理；（B）若仅"科研 / BCI / 神经反馈"用途 → 有二类空间但市场价值大幅缩水'],
    ['风险等级参照', '2020 版《医疗器械分类目录》中"经颅电刺激仪"多按 三类；参考已批品种（如某国产 rTMS/tDCS 治疗仪）多为三类证'],
    ['临床评价路径', '不在豁免目录内，必须开展中国境内多中心临床试验（GCP）'],
    ['本报告基准假设', '按"三类 + 治疗适应症"推进，与对方"走临床"策略一致；若最终降级二类，可参考 Enobio 的成本结构'],
]
add_table(doc, ['项目', '内容'], starstim_rows, col_widths=[4, 12])

doc.add_page_break()

# ==== 三、法规环境 ====
add_heading(doc, '三、中国 NMPA 法规环境与产品分类判定', level=1)

add_heading(doc, '3.1 医疗器械分类界定', level=2)
add_para(doc,
    '根据《医疗器械监督管理条例》（2021 修订，国务院令第 739 号）与《医疗器械分类目录》（2017 公布、'
    '持续更新）：一类备案（低风险）、二类地方局注册（中风险）、三类国家局注册（高风险）。'
)
add_para(doc,
    '经查询目录及历史注册数据：'
)
add_bullet(doc, 'EEG（脑电图机、脑电采集与分析系统）：二类，分类编码 07-03-04 / 06-09。')
add_bullet(doc, '经颅电刺激类（tDCS / tACS / 经颅微电流刺激）：若声称治疗神经/精神疾病，按 三类，编码 09-03；若声称"心理调节/减压"等消费级用途，另有二类甚至非医疗器械通道，但不适用于本次拟报路径。')

add_heading(doc, '3.2 进口与"国产"两种路径的本质差异', level=2)
pathway_rows = [
    ['申请人主体', '境外原厂 / 香港代理', '中国境内制造企业（境内注册公司 + 境内生产场地）'],
    ['注册体系', '进口注册（NMPA 国家局）', '境内注册（二类地方局 / 三类国家局）'],
    ['体系核查', '境外体系核查（可能 remote，二三类高风险抽查）', '境内 GMP 核查（国产必查）'],
    ['技术转让/授权', '无需，直接以境外主体申报', '必须有厂家授权 / 技术转让协议'],
    ['注册周期', '二类约 12–18 月，三类 24–36 月+', '二类 12–18 月，三类 30–48 月（含体系搭建）'],
    ['注册费用（官费）', '二类 21.09 万，三类 30.88 万（境外首次）', '二类 免收部分官费（地方局），三类 15.36 万（境内首次）'],
    ['集采/市场准入', '进口证在部分省份集采不占优、招标常被"国产优先"扣分', '国产证在集采、大三甲招标、DRG/DIP 报销中优势明显'],
    ['整体推荐度（本项目）', '一般', '推荐——这也是委托方明确要求的路径'],
]
add_table(doc, ['维度', '进口路径', '国产路径（推荐）'], pathway_rows,
          col_widths=[3, 6, 7])

add_heading(doc, '3.3 免于临床试验目录（豁免临床）适用性分析', level=2)
add_para(doc,
    '2023 年 9 月 NMPA 修订发布的《免于临床评价医疗器械目录》与《医疗器械临床评价技术指导原则》'
    '（2021 年 43 号通告）规定：目录内产品可通过"同品种比对 + 非临床数据"完成临床评价，'
    '无需开展中国境内临床试验。'
)
add_bullet(doc, 'Enobio（EEG）：目录第 07 类"医用电子仪器设备"下"脑电图机"（分类编码 07-03-04）在豁免范围内。')
add_bullet(doc, 'Starstim（tES）：目录中"经颅电刺激"类未列入豁免；即便部分低电流"经颅微电流刺激"曾有二类豁免案例，凡带"治疗适应症"的均需临床数据支持。')
add_para(doc,
    '结论：Enobio 走"同品种比对"路径合规且高效；Starstim 无豁免可能，必须做临床试验。这也是'
    '"两个产品价格差数量级"的根本原因，也印证对方代表分别报价的合理性。'
)

doc.add_page_break()

# ==== 四、国产化落地方案 ====
add_heading(doc, '四、"以国产形式申报"的落地方案', level=1)

add_heading(doc, '4.1 三种国产化模式对比', level=2)
mode_rows = [
    ['模式 A：自建境内工厂', '设立 WFOE 或合资公司，购置产线，境内完成整机组装、检测', '真国产、完全自主', '设立成本 300–800 万；周期 12–18 个月', '一次性投入大，但适合长期战略布局'],
    ['模式 B：中国 OEM 代工 + 委托生产', '与国内已获相关许可的医疗器械厂签订委托协议，由其代工，注册人自行申报', '轻资产、6–12 个月即可具备申报条件', '需支付代工费/管理费；核心技术依赖厂家', '推荐——性价比最高的过渡方案'],
    ['模式 C：贴牌授权（License + Rebranding）', '中国厂家取得境外授权，以自有品牌申报（非注册人制度）', '最快、投入最少', '境外品牌几乎归零、后续升级依赖境外方', '不满足"引进两款产品"的战略目的，仅作备选'],
]
add_table(doc, ['模式', '核心机制', '优点', '成本/周期', '本项目建议'], mode_rows,
          col_widths=[3, 3.5, 3, 3.5, 3])

add_heading(doc, '4.2 生产场地与体系建设要求', level=2)
add_bullet(doc, '生产许可证：二类需省级药监局《第二类医疗器械生产许可证》；三类另需三类生产许可（更严格的体系核查）。')
add_bullet(doc, 'ISO 13485 / YY/T 0287 质量管理体系：全面覆盖设计开发、采购、生产、检验、售后。')
add_bullet(doc, '洁净车间：EEG/tES 属"非无菌电子医疗器械"，一般不需要 10 万级洁净，普通电子装配车间 + ESD 控制即可，但需静电、温湿度、防尘等基础要求。')
add_bullet(doc, '关键工序自制比例：NMPA 对"国产"实质性生产有隐性要求，通常至少完成整机组装、软件烧录、终检、出厂放行，若关键 PCBA 全部进口需在技术审评环节补充说明。')

add_heading(doc, '4.3 关键前提：与西班牙方的技术转让/授权协议', level=2)
add_para(doc,
    '"国产化申报"的法律基石是"注册申请人=境内生产企业"，因此必须与 Neuroelectrics 签署具备如下'
    '关键条款的协议，否则无法立项：'
)
add_bullet(doc, '技术转让（Technology Transfer）：设计文件、结构图纸、PCB 原理图、软件源代码/固件、检验规范、生产工艺文件。')
add_bullet(doc, '生产授权（Manufacturing License）：授权中国主体在中国境内合法生产、销售、贴上中国主体商标。')
add_bullet(doc, '知识产权与专利许可：明确 IP 归属、独家/非独家、地域范围（建议至少大中华区独家）、期限。')
add_bullet(doc, '临床数据可用性：Neuroelectrics 已有 CE 临床数据、发表论文、上市后监测数据的"可引用授权"，直接决定同品种比对的可行性与工作量。')
add_bullet(doc, '售后与升级：固件升级、备件供应、故障响应 SLA。')
add_para(doc,
    '这些条款的谈判成本（法务、专利检索、境外尽调）并未包含在对方 100–150 万报价中，'
    '预算需额外预留 30–80 万法务/谈判费用。'
)

doc.add_page_break()

# ==== 五、注册路径与临床策略 ====
add_heading(doc, '五、注册路径与临床策略详解', level=1)

add_heading(doc, '5.1 Enobio（EEG）：对标同品种、豁免临床路径', level=2)
add_para(doc, '推荐流程（8 大步骤）：', bold=True, indent_cn=False)
enobio_steps = [
    ['1', '产品分类界定', '向药监局分类界定办公室提交《分类界定申请》，明确二类、编码 07-03-04', '1–2 月'],
    ['2', '技术要求与检验', '编制产品技术要求（YY 0505 电磁兼容、GB 9706.1 电气安全、YY 9706.226 EEG 专用标准）→ 送国家或省级医疗器械检验所全项检测', '4–6 月'],
    ['3', '同品种比对（临床评价）', '选择 1–2 款已上市 EEG（如迈瑞 mNC-A/新松），逐项比对基本原理、结构组成、性能指标、适用范围、安全有效性；差异点用非临床数据支持', '2–3 月'],
    ['4', '质量体系建立', '按 ISO 13485 / YY/T 0287 建立文件化体系，完成内审、管审', '3–6 月（与检验并行）'],
    ['5', '生产许可申报', '省药监局申报《第二类医疗器械生产许可证》，现场核查', '2–4 月'],
    ['6', '注册申报（境内二类）', '向省级药监局提交注册资料（技术要求、检验报告、临床评价报告、说明书、体系文件等）', '1 月递交'],
    ['7', '技术审评 + 补正', '省局审评中心受理→审评→发补→补正→再审', '6–10 月'],
    ['8', '注册证发放', '通过后 20 工作日内发证，有效期 5 年', '1 月'],
]
add_table(doc, ['步骤', '事项', '关键要点', '时间'], enobio_steps,
          col_widths=[1.2, 3.5, 8.8, 2.5])
add_para(doc, '合计周期：12–18 个月；合计费用（下节详拆）：100–180 万人民币。',
         bold=True, indent_cn=False, color=(192, 0, 0))

add_heading(doc, '5.2 Starstim（tES）：必须做临床的路径', level=2)
add_para(doc, '推荐流程（10 大步骤）：', bold=True, indent_cn=False)
starstim_steps = [
    ['1', '产品分类界定 + Pre-sub', '向 CMDE 申请 Pre-submission 沟通，明确三类、编码 09-03，确认临床方案框架', '2–3 月'],
    ['2', '产品技术要求与检验', 'GB 9706.1、YY 9706.240（如适用）、生物相容性（GB/T 16886 系列，皮肤接触电极）、软件、网络安全（YY/T 0664 / YY/T 1843）等', '6–9 月'],
    ['3', '动物实验（如需要）', 'tES 若已有充分文献支持，可豁免；否则需补做，通常猪/大鼠皮肤刺激 + 电流耐受', '3–6 月（可选）'],
    ['4', '临床方案设计 + 伦理', '选择适应症（推荐"卒中后运动/认知康复"或"重度抑郁"为首选，已有较多国际证据支持）；多中心 RCT；样本量 100–200 例；随访 3–6 月', '3–6 月'],
    ['5', '临床试验备案（三类需备案）', 'NMPA 器审中心（CMDE）备案，60 工作日内如无异议即可启动', '3 月'],
    ['6', '多中心临床执行', '3–5 家 GCP 认证三甲医院同步入组，含伦理复审、CRC 驻点、SDV 监查、SAE 报告', '18–24 月'],
    ['7', '数据锁定与统计分析', 'CRO 或独立统计单位完成 SAP、盲态审核、Database Lock、SAS 分析、总结报告', '3–4 月'],
    ['8', '注册申报（三类）', '向 NMPA 国家局提交完整注册资料（含 CER、临床试验总结报告）', '1 月递交'],
    ['9', 'CMDE 技术审评 + 体系核查 + 发补', '一般 1–2 轮发补，境内生产企业需现场体系核查', '12–18 月'],
    ['10', '注册证发放', '有效期 5 年', '1–2 月'],
]
add_table(doc, ['步骤', '事项', '关键要点', '时间'], starstim_steps,
          col_widths=[1.2, 3.5, 8.8, 2.5])
add_para(doc, '合计周期：30–48 个月；合计费用（下节详拆）：550–950 万人民币。',
         bold=True, indent_cn=False, color=(192, 0, 0))

add_heading(doc, '5.3 同步 vs 依次策略的定量对比', level=2)
strategy_rows = [
    ['同步策略（Parallel）', '两个产品并行推进注册与临床', '整体周期以 Starstim 为准（36 月）；两个产品几乎同时上市', '+150 万设备成本（多一套 15–20 台）+ 30–50 万额外 CRC / 监查', '首年市场机会最大化；总成本最高'],
    ['半同步（Enobio 先 + Starstim 立即启动）', 'Enobio 走豁免临床，Starstim 同步启动临床', 'Enobio 12–18 月拿证（先卖 EEG 造血）；Starstim 30–48 月拿证', '节省一部分同步成本 60–100 万', '折中方案，推荐——现金流与市场兼顾'],
    ['依次策略（Sequential）', '先做 Enobio 拿证，再启动 Starstim', 'Enobio 12–18 月拿证；Starstim 之后 30–48 月（合计 42–66 月）', '节省 150 万设备 + 部分人员成本；但延误上市 12–18 月', '总花费最省，但市场机会成本极高'],
]
add_table(doc, ['策略', '推进方式', '总周期', '增量/节省', '综合评价'], strategy_rows,
          col_widths=[3, 3, 3, 3.5, 3.5])

doc.add_page_break()

# ==== 六、时间进度评估 ====
add_heading(doc, '六、时间进度评估', level=1)

add_heading(doc, '6.1 关键里程碑与关键路径', level=2)
add_para(doc,
    'Starstim 的关键路径是"临床试验执行 18–24 月 + 技术审评 12–18 月"，两者累计 30–42 月，'
    '几乎决定整个项目的最短工期。任何前置工作（体系搭建、检验、伦理）都可以在这段时间内并行完成，'
    '因此优化重点应放在"临床提前启动"与"CMDE 沟通前置"上：'
)
add_bullet(doc, '尽早（分类界定后 3 个月内）向 CMDE 申请 Pre-submission 沟通，可为后续注册节省 3–6 个月发补时间。')
add_bullet(doc, '临床方案与产品技术要求同步定稿，避免临床数据不匹配技术要求导致返工。')
add_bullet(doc, '选择适应症时，优先"已有 CE 数据 + 中国指南推荐"的双重支持适应症（如卒中后运动康复），避免"抑郁症"等审评严格且长期争议的领域。')

add_heading(doc, '6.2 三种策略时间对比表', level=2)
time_rows = [
    ['分类界定 / Pre-sub', '同步：0–3 月', '半同步：0–3 月', '依次：0–3 月'],
    ['Enobio 拿证', '同步：12–18 月', '半同步：12–18 月', '依次：12–18 月'],
    ['Starstim 临床启动', '同步：3–6 月', '半同步：3–6 月', '依次：15–20 月（Enobio 拿证后启动）'],
    ['Starstim 拿证', '同步：30–48 月', '半同步：30–48 月', '依次：45–66 月'],
    ['两证齐全（全部上市）', '同步：30–48 月', '半同步：30–48 月', '依次：45–66 月'],
    ['市场先机（首证上市）', '同步：12–18 月（Enobio）', '半同步：12–18 月', '依次：12–18 月'],
]
add_table(doc, ['里程碑', '同步策略', '半同步策略', '依次策略'], time_rows,
          col_widths=[4, 4, 4, 4])

add_para(doc,
    '半同步策略在"首证时间"上与同步策略无差异（都是 12–18 月拿到 EEG 证），但可以省下大部分'
    '"同时启动"引发的重复投入，是本报告最推荐的策略。'
)

doc.add_page_break()

# ==== 七、费用测算 ====
add_heading(doc, '七、费用测算与拆解', level=1)

add_heading(doc, '7.1 Enobio 单独注册费用（无临床）', level=2)
enobio_cost = [
    ['产品技术要求编制', '5–10 万', '含标准检索、技术要求文本、验证报告；可自行完成或外包'],
    ['注册检验（型式试验）', '20–35 万', 'GB 9706.1 + EMC + EEG 专用标准 + 软件 + 网络安全；国家医疗器械检验所或有资质的第三方'],
    ['同品种比对（临床评价报告 CER）', '8–15 万', '外包咨询公司或自行完成；文献综述 + 差异分析'],
    ['质量体系建立（ISO 13485）', '15–30 万', '文件体系 + 培训 + 内审 + 认证辅导；首次搭建'],
    ['生产许可申报', '3–5 万', '资料准备 + 现场核查配合'],
    ['注册申报官费（省级二类）', '0（多省免收）', '各省政策不同，多数二类境内注册免官费'],
    ['注册代理/咨询费', '30–50 万', '资料撰写、发补应对、审评沟通、全流程管理'],
    ['其他（法务、翻译、差旅）', '10–20 万', '技术转让协议翻译公证、差旅、专家论证'],
    ['小 计', '90–165 万', '与对方报价 100–150 万高度吻合'],
]
add_table(doc, ['项目', '金额（人民币）', '说明'], enobio_cost,
          col_widths=[5, 3.5, 7.5])

add_heading(doc, '7.2 Starstim 带临床注册全成本', level=2)
add_para(doc, '（按 3 中心、150 例、随访 3 个月的常规三类神经调控临床估算）', bold=True, indent_cn=False)
starstim_cost = [
    ['A. 注册前置成本', '', ''],
    ['产品技术要求 + 型式检验', '35–60 万', '三类要求更多测试（含生物相容性、软件、网络安全）'],
    ['动物实验（如需）', '30–60 万', '若能引用 CE/FDA 已有数据可省'],
    ['质量体系（三类要求）', '30–60 万', 'ISO 13485 + 三类生产许可辅导'],
    ['临床方案设计 + 伦理准备', '15–30 万', 'CRO 或注册顾问'],
    ['B. 临床试验成本（核心）', '', ''],
    ['CRO 项目管理费', '80–150 万', '3 中心 × 12–18 月，含 CRA 监查'],
    ['研究者费用（PI/Co-I）', '90–180 万', '3 中心 × 30–60 万/中心（含PI酬金、协作费）'],
    ['受试者费用（补偿+检查）', '90–180 万', '150 例 × 6000–12000 元（含 MRI/量表评估/交通/耽误费）'],
    ['CRC / 数据管理', '40–80 万', '每中心 CRC 或电子源数据搭建'],
    ['统计分析 / SAP / DSMB', '20–40 万', '独立统计单位'],
    ['伦理 / 保险 / SAE 报告', '15–30 万', '临床试验保险约 3–8 万/中心'],
    ['C. 临床设备与耗材', '', ''],
    ['设备（Starstim 15–20 台）', '150–200 万', '按 8–10 万/台成本价（对方口径），实际设备 BOM 可能更低'],
    ['配套 Enobio（若需要）', '80–120 万', '若临床同时采集脑电对比'],
    ['耗材（电极、凝胶、贴片）', '20–35 万', '3 中心 × 150 例的实际消耗'],
    ['D. 注册申报环节', '', ''],
    ['注册申报官费（三类境内）', '15.36 万', 'NMPA 明码收费'],
    ['注册代理 / 咨询 / 发补应对', '80–150 万', '三类审评复杂，代理费显著高于二类'],
    ['E. 隐性成本', '', ''],
    ['技术转让 / 授权费', '50–200 万', '与 Neuroelectrics 谈判结果，可能一次性或分期'],
    ['境外体系核查支持（若涉及原厂参与）', '10–30 万', ''],
    ['小 计（不含隐性）', '740–1315 万', ''],
    ['合 计（含隐性）', '800–1550 万', '典型区间 900–1200 万'],
]
add_table(doc, ['项目', '金额（人民币）', '说明'], starstim_cost,
          col_widths=[5, 3.5, 7.5])

add_heading(doc, '7.3 与西班牙代表报价的逐项比对', level=2)
compare_rows = [
    ['注册（无临床）', '100–150 万', '90–165 万', '一致', '对方报价合理，落在市场标准区间下限'],
    ['设备成本（15–20 台）', '150 万', '150–200 万', '基本一致', '按 8–10 万/台成本价，是原厂对内部关联方转移价格；市面公开售价为其 2–3 倍'],
    ['耗材（临床用）', '23 万', '20–35 万', '一致', '合理'],
    ['医生费用（含 PI 等）', '250 万', '400–800 万（常规） / 250–400 万（关系型压价）', '偏低', '如靠关系压价 + 论文归属激励可实现，但存在合规风险；见 8.3 详析'],
    ['同步 +150 万', '150 万（多一套设备）', '180–220 万（含 CRC 与监查额外费用）', '偏低', '仅覆盖设备，未含增量人力'],
    ['依次省 150 万', '省一套设备', '省 150 万但延误 12–18 月上市', '结论合理', '但机会成本通常远高于 150 万'],
]
add_table(doc, ['分项', '对方报价', '本报告测算', '偏差', '备注'], compare_rows,
          col_widths=[3, 3, 3.5, 2.5, 5])

doc.add_page_break()

# ==== 八、对方报价合理性判定 ====
add_heading(doc, '八、对方报价合理性判定', level=1)

add_heading(doc, '8.1 "100–150 万无临床注册费"合理性', level=2)
add_para(doc,
    '判定：合理，且属于市场"熟练代理"的下限报价。'
    '国内二类境内注册的市场行情：完全外包（含检验、CER、体系、代理）'
    '通常 120–200 万；若企业自建部分能力，可压到 90–130 万。'
    '对方给到 100–150 万是通行价格，无异常。',
    color=(0, 112, 74),
)

add_heading(doc, '8.2 "150 万设备成本、23 万耗材"合理性', level=2)
add_para(doc,
    '判定：合理，但需注意两点。',
    color=(0, 112, 74),
)
add_bullet(doc, '150 万设备 ÷ 20 台 = 7.5 万/台，与 Starstim 32 通道市场售价 25–35 万（研究型）相比，实为 BOM/成本价，说明西班牙原厂愿意以"内部关联方转移价"支持临床。这是合理商业安排，但要在合同中锁死"临床用设备仅供 xx 项目、不得二次销售"以规避税务和市场秩序问题。')
add_bullet(doc, '23 万耗材对于 3 中心 150 例的用量偏保守（一次性电极片、导电膏、贴片 + 备品）。若考虑设备故障备件损耗，实际可能在 25–35 万区间。')

add_heading(doc, '8.3 "250 万医生费用"合理性（重点关注）', level=2)
add_para(doc,
    '判定：偏低，需谨慎评估其真实覆盖范围。',
    color=(192, 0, 0),
)
add_para(doc, '合规三类神经调控多中心临床的"医生 + 临床"支出结构如下：')
doctor_rows = [
    ['PI 主要研究者酬金', '3 中心 × 30–50 万', '90–150 万', '按项目分期发放，含论文写作、启动会、答辩'],
    ['Co-Investigator / 协作团队', '3 中心 × 15–30 万', '45–90 万', '含康复师、护理团队、影像医师'],
    ['受试者补偿 + 检查费', '150 例 × 6000–12000 元', '90–180 万', '含 MRI/EEG/量表评估等衍生检查'],
    ['CRC 驻点', '3 CRC × 12 月 × 1.5 万', '54 万', '临床协调员，SDV 前的数据整理'],
    ['伦理审查费', '3 中心 × 3–5 万', '9–15 万', ''],
    ['CRO 监查（CRA/PM）', '', '80–150 万', '与 CRO 项目管理费部分重合'],
    ['统计 / 数据管理', '', '20–40 万', ''],
    ['临床试验保险', '', '15–30 万', ''],
    ['合计（常规合规成本）', '', '400–700 万', '含 CRO 部分'],
]
add_table(doc, ['项目', '基准', '金额', '说明'], doctor_rows,
          col_widths=[4, 4, 3, 5])

add_para(doc,
    '对方报出的 250 万如果仅指"PI + Co-Investigator + 受试者补偿"三项主要花在'
    '"人"上的费用，且通过下述途径实现压价，是可能的：',
)
add_bullet(doc, 'PI 与集团有历史合作或利益共同体（"关系型"），费用低于市场价。')
add_bullet(doc, '"论文归属医生"作为非现金激励，替代部分酬金；但注意 GCP 规范下论文署名与费用支付需分离、有独立性说明，不能作为"以论文换减免"。')
add_bullet(doc, '选择 2 家而非 3 家中心（多中心最低 2 家即可满足要求）。')
add_bullet(doc, '样本量控制在 100 例而非 150 例（需在方案与统计假设中论证 ≥80% power）。')
add_para(doc,
    '但即便如此，本报告仍强烈提醒：CRO 监查、数据管理、统计、伦理、保险这些"非医生但必需"的费用'
    '（合计 150–250 万）不能省，也不在"医生费用"科目内。如果对方 250 万是"总临床费用"'
    '（含 CRO 等），则严重低估；如果是"仅医生 + 受试者"科目，则相对合理但仍需谈判确认。',
    bold=True,
)

add_heading(doc, '8.4 "同步 +150 万 / 依次省设备"策略合理性', level=2)
add_para(doc,
    '判定：方向合理，量化偏乐观。',
    color=(0, 112, 74),
)
add_bullet(doc, '同步策略确实需要多备 15–20 台设备（BOM 价 ~150 万），但同步执行还会带来'
                '"额外 CRC、额外监查、额外耗材、额外场地"的增量 30–70 万，不能忽略。')
add_bullet(doc, '依次策略节省 150 万设备是显性的，但延迟 12–18 个月上市的机会成本，'
                '按 Starstim 单价 25 万、年销 100 台估算，就是 2500 万/年的销售延迟，'
                '远远超过节省的 150 万；除非现金流紧张，否则不推荐依次。')
add_para(doc,
    '综上，"同步"或"半同步"是从财务净现值角度最优的策略。',
    bold=True,
)

doc.add_page_break()

# ==== 九、风险清单 ====
add_heading(doc, '九、风险清单与不确定性提示', level=1)

risk_rows = [
    ['R1', '分类界定风险', '中', 'Starstim 最终被判定为三类是主流预期；但若声称"睡眠改善、焦虑辅助"等适应症，可能被要求补充疗效证据甚至转为消费级/非医疗器械。', '尽早向 CMDE 提交 Pre-submission，锁定分类与适应症'],
    ['R2', '临床数据不接受', '中', 'CMDE 对"引用境外数据"审查趋严，境外多中心 RCT 数据仅作参考，仍需中国境内主临床。', '按本报告方案设计中国境内 RCT，境外数据作补充证据'],
    ['R3', '技术转让协议谈判失败', '低-中', 'Neuroelectrics 若不愿转让固件源码或关键工艺，会影响国产化实质性判定。', '合同前置尽调；预留 30–80 万法务预算'],
    ['R4', '生产体系核查未通过', '中', '首次搭建 ISO 13485 与三类生产许可，现场核查一次通过率低于 50%。', '聘请有辅导资质的咨询公司；预留 3–6 月整改期'],
    ['R5', '临床 SAE / 期中分析失败', '中-高', 'tES 类产品在焦虑、抑郁、认知等适应症上历史阴性结果比例较高。', '选择"卒中运动康复"等高证据强度适应症；设置 DSMB 中期分析'],
    ['R6', '医生费用估算失真', '中', '若"关系"路径无法兑现，或 GCP 合规不允许论文换酬金，250 万预算不足以完成合规临床。', '按合规区间（400–800 万）预留、以关系压价为上侧惊喜而非基线'],
    ['R7', '外汇与关税', '低', '设备从西班牙 BOM 入华涉及关税、增值税，或按内部关联方转让定价审查。', '税务顾问先行方案设计'],
    ['R8', '汇率与预算通胀', '低', '3–4 年周期内人民币/欧元汇率、临床费用通胀', '预算预留 10% 缓冲'],
    ['R9', 'DRG/DIP 医保准入', '中', '即便拿证，医保准入决定销售规模', '同步与目标医院、地方医保沟通，争取新技术加成'],
    ['R10', '竞品加速', '中', '国内 rTMS/tDCS/tACS 品牌（如脑陆、诺辉、湖南埃普特）在 2024–2025 年密集获批', '争取半同步策略，尽早首证'],
]
add_table(doc, ['编号', '风险', '概率', '影响与情景', '缓释措施'], risk_rows,
          col_widths=[1.2, 3, 1.5, 6.5, 4])

doc.add_page_break()

# ==== 十、结论与建议 ====
add_heading(doc, '十、结论与行动建议', level=1)

add_heading(doc, '10.1 核心判断', level=2)
add_bullet(doc, '对方代表对"注册路径"层面的理解基本正确，报价也在市场合理区间——尤其是 100–150 万无临床注册费与设备/耗材成本部分。')
add_bullet(doc, '250 万医生费用是最需要澄清的口径：若仅指"PI + 受试者"人费，可接受但偏紧；若号称覆盖"整个临床"，则严重低估。')
add_bullet(doc, '"同步 +150 万"是理性化描述，但缺 30–70 万增量执行成本；"依次省 150 万"仅在极度现金流紧张时才合理。')
add_bullet(doc, '"国产化"本身的隐性成本（体系搭建、技术授权、税务筹划）是对方报价中最大的缺失项，需自行预留 300–800 万。')

add_heading(doc, '10.2 推荐总体预算区间', level=2)
budget_rows = [
    ['Enobio 注册（无临床，国产化路径）', '100–180 万', '含 90–130 万注册 + 20–50 万体系与生产许可 + 10–30 万授权与法务'],
    ['Starstim 注册（含临床，国产化路径）', '900–1200 万（典型）', '含 500–800 万临床全成本 + 150–200 万设备 + 20–35 万耗材 + 15.36 万官费 + 80–150 万代理 + 50–200 万授权隐性'],
    ['同步策略总预算', '1000–1380 万', '两个产品同步推进'],
    ['半同步策略总预算', '1000–1350 万（推荐）', 'Enobio 先拿证造血，Starstim 同步启动临床'],
    ['依次策略总预算', '900–1250 万', '省 100–150 万但延误上市，NPV 不划算'],
]
add_table(doc, ['策略/项目', '预算区间（人民币）', '备注'], budget_rows,
          col_widths=[6, 4, 6])

add_heading(doc, '10.3 时间预期', level=2)
add_bullet(doc, 'Enobio 拿证：项目立项后 12–18 个月')
add_bullet(doc, 'Starstim 拿证：项目立项后 30–48 个月（同步或半同步策略）')
add_bullet(doc, '首证上市（Enobio）：12–18 个月即可产生现金流')
add_bullet(doc, '完整两证：36 个月为进取目标、48 个月为稳健目标')

add_heading(doc, '10.4 谈判与合同层面的关键动作', level=2)
add_bullet(doc, '（合同条款）向西班牙原厂锁定：技术转让范围、生产授权地域独家性、论文与临床数据可引用性、售后 SLA、备件价格年增幅上限。', bold_prefix='① ')
add_bullet(doc, '（费用口径）要求对方以书面形式确认 250 万医生费用的"具体覆盖科目"，并明确 CRO、监查、伦理、保险是否额外。', bold_prefix='② ')
add_bullet(doc, '（分类界定）第 1 个月即启动 CMDE Pre-sub，锁定 Starstim 三类与适应症选择。', bold_prefix='③ ')
add_bullet(doc, '（体系落地）确定国产化模式（自建 vs OEM 代工），签署代工/生产协议，避免体系空转。', bold_prefix='④ ')
add_bullet(doc, '（现金流规划）按半同步策略滚动预算，Enobio 拿证后销售现金流可以对冲 Starstim 临床费用高峰。', bold_prefix='⑤ ')
add_bullet(doc, '（风险准备金）在总预算之上再预留 10%（100–140 万）作为发补、重检、SAE 处置准备金。', bold_prefix='⑥ ')

add_heading(doc, '10.5 决策建议', level=2)
add_para(doc,
    '综合考虑法规可行性、时间机会成本与总投入，推荐执行【半同步策略】：',
    bold=True,
)
add_bullet(doc, '立项 0–3 月：完成分类界定、Pre-sub 沟通、与西班牙方签署技术转让/授权协议。')
add_bullet(doc, '立项 3–6 月：Enobio 与 Starstim 同步启动体系搭建、Starstim 启动临床方案与伦理。')
add_bullet(doc, '立项 6–18 月：Enobio 完成检验、注册申报、拿证；Starstim 完成临床首例入组、持续入组。')
add_bullet(doc, '立项 18–30 月：Enobio 首年销售；Starstim 完成入组、数据锁定、总结报告。')
add_bullet(doc, '立项 30–48 月：Starstim 完成 CMDE 审评、拿证、上市。')
add_para(doc,
    '在此策略下，总预算区间 1000–1350 万人民币是合理规划基准；对方报价（不含"国产化隐性成本"'
    '与"完整临床合规费用"）大约位于 550–750 万区间，与本报告差 300–600 万，主要差在'
    '"体系搭建/授权/CRO 与合规"三块。这笔差额不是对方"骗人"，而是对方口径下不含的项目，'
    '需要在决策前预算清楚，避免后续被动。',
)

doc.add_page_break()

# ==== 附录 ====
add_heading(doc, '附录 A：假设与数据来源说明', level=1)
add_bullet(doc, '医疗器械分类依据：《医疗器械分类目录》（国家药监局 2017 年公告，含 2023 年动态调整）。')
add_bullet(doc, '临床评价路径：《医疗器械临床评价技术指导原则》（2021 年 43 号通告）、《免于临床评价医疗器械目录》（2023 年 9 月修订）。')
add_bullet(doc, '注册官费：《国家药品监督管理局医疗器械注册收费标准》（现行）。')
add_bullet(doc, '临床试验费用：基于 2022–2025 年国内三类神经调控/康复类多中心 RCT 项目公开成本经验区间，未考虑区域与医院差异。')
add_bullet(doc, '设备与耗材成本：以对方代表提供口径为基线，市场公开售价为参照，未做独立 BOM 核算。')
add_bullet(doc, '本报告不构成法律/税务意见，最终决策请结合专业注册顾问、律师、税务师复核。')

add_heading(doc, '附录 B：术语表', level=1)
term_rows = [
    ['NMPA', '国家药品监督管理局（National Medical Products Administration）'],
    ['CMDE', '国家药监局医疗器械技术审评中心（Center for Medical Device Evaluation）'],
    ['GCP', 'Good Clinical Practice，药物/器械临床试验质量管理规范'],
    ['ISO 13485', '医疗器械质量管理体系国际标准'],
    ['YY/T', '中国医药行业推荐性技术标准'],
    ['GB 9706.1', '医用电气设备通用安全标准（2020 修订）'],
    ['CER', 'Clinical Evaluation Report，临床评价报告'],
    ['Pre-sub', 'Pre-submission，与 CMDE 的正式注册前沟通'],
    ['SAE', 'Serious Adverse Event，严重不良事件'],
    ['DSMB', 'Data Safety Monitoring Board，数据安全监察委员会'],
    ['CRA', 'Clinical Research Associate，临床监查员'],
    ['CRC', 'Clinical Research Coordinator，临床协调员'],
    ['SDV', 'Source Document Verification，原始文档核查'],
    ['DRG/DIP', '按疾病诊断相关分组付费/按病种分值付费'],
    ['tES / tDCS / tACS / tRNS', '经颅电刺激/经颅直流电刺激/经颅交流电刺激/经颅随机噪声刺激'],
    ['EEG', 'Electroencephalography，脑电图/脑电采集'],
]
add_table(doc, ['术语', '含义'], term_rows, col_widths=[4, 12])

# 末尾说明
add_para(doc, '')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('—— 报告完 ——')
set_run_font(r, size=11, color=(89, 89, 89))

# 保存
out_path = '/home/user/sugz/scratchpad/Neuroelectrics_引进国产化注册可行性评估报告.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
