# -*- coding: utf-8 -*-
"""
Flow FL-100 全部临床前与临床研究清单表
——用于 Starstim 中国注册工作分解
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_run_font(run, name='微软雅黑', size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rpr = r.rPr
    if rpr is None:
        rpr = OxmlElement('w:rPr')
        r.insert(0, rpr)
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def add_heading(doc, text, level=1, color=None):
    styles = {
        1: (16, True, (31, 73, 125)),
        2: (13.5, True, (54, 95, 145)),
        3: (12, True, (68, 114, 148)),
    }
    size, bold, default_color = styles.get(level, (11, True, (0, 0, 0)))
    if color is None:
        color = default_color
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_para(doc, text, bold=False, size=10.5, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_table(doc, headers, rows, col_widths=None, header_color='1F497D', font_size=9):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=9.5, bold=True, color=(255, 255, 255))
        set_cell_bg(hdr[i], header_color)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ''
            p = cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            set_run_font(run, size=font_size)
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


# ============ 生成文档 ============
doc = Document()

# 横向页面
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    # 横向
    new_width, new_height = section.page_height, section.page_width
    section.orientation = 1  # WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ==== 标题 ====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Flow FL-100 全部临床前 + 临床研究清单表')
set_run_font(r, size=18, bold=True, color=(31, 73, 125))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('（用于 Starstim 中国 NMPA 三类注册的工作分解清单）')
set_run_font(r, size=12, color=(89, 89, 89))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('来源：FDA SSED, PMA P230024（Flow FL-100，2025-12-08 批准）')
set_run_font(r, size=10, color=(89, 89, 89))

# ==== 一、临床前研究清单 ====
add_heading(doc, '一、临床前研究清单（Pre-clinical Studies）——共 5 项', level=1)

preclinical_rows = [
    # 编号, 研究项目, Flow 具体做法, 判定/结果, 依据标准, 中国对应标准/要求, 预计时间, 预计费用
    ['P-01', '生物相容性评价',
     '按 ISO 10993-1 风险管理路径评价\n接触部位：完整皮肤\n接触时长：单次≤30 min，累计>24h（分散数月）\n材料：ABS 塑料 + 医用硅胶 + 纤维素海绵',
     '风险低，通过',
     'ISO 10993-1\nFDA Attachment G "Biocompatibility of Certain Devices in Contact with Intact Skin"',
     'GB/T 16886.1 生物学评价与试验总则\nGB/T 16886.5 细胞毒性\nGB/T 16886.10 皮肤致敏\nGB/T 16886.23 皮内反应\n（可引用境外数据减少国内测试）',
     '2–4 月',
     '15–30 万'],

    ['P-02', '电磁兼容与电气安全',
     '按 3 项国际标准测试\n（含家用医疗设备并列标准）',
     '合规',
     'IEC 60601-1 通用安全\nIEC 60601-1-2 EMC 并列\nIEC 60601-1-11 家用医疗设备并列',
     'GB 9706.1-2020 通用安全\nYY 9706.102-2021 EMC 并列\nYY 9706.111 家用医疗设备并列\nYY 9706.210-2021 神经肌肉刺激器专用要求',
     '4–6 月',
     '25–45 万'],

    ['P-03', '软件生存周期与网络安全',
     'App 端（患者）\nCPP 医生端\n服务器后端\n软件安全等级：Major level of concern（三类要求）',
     '文档充分\n满足三类要求',
     'FDA Software Guidance\nFDA Cybersecurity Guidance\nSaMD (Software as Medical Device)',
     'YY/T 0664-2020 医疗器械软件生存周期\nYY/T 1843-2022 医疗器械网络安全\nNMPA 软件注册审查指导原则 2022\nNMPA 网络安全注册审查指导原则 2022',
     '3–6 月（与研发并行）',
     '20–40 万'],

    ['P-04', '有效期验证（Shelf Life）',
     '实测 + 加速老化推算\n头戴装置 + 一次性电极片',
     '3 年',
     'ASTM F1980 加速老化\nISO 11607 包装',
     'YY/T 0681 系列 无菌医疗器械包装试验（可参考）\n注册需附加速老化 + 实时老化数据',
     '6–12 月（加速老化）',
     '10–20 万'],

    ['P-05', '性能验证（Performance Testing）',
     '每台设备出厂验证输出电流 2.0 ± 0.05 mA',
     '合规\n生产环节 100% 检验',
     '产品技术要求自定义',
     '产品技术要求（YY/T 组织自编）需覆盖：\n① 电流精度\n② 多通道波形一致性\n③ 总电荷密度安全阈值\n④ 皮肤接触阻抗监控\n⑤ 蓝牙通信稳定性',
     '3–6 月',
     '10–25 万'],
]
add_table(doc,
    ['编号', '研究项目', 'Flow FL-100 具体做法', 'Flow 判定/结果', 'FDA 依据标准',
     '中国对应标准/NMPA 要求', 'Starstim 预计时间', 'Starstim 预计费用'],
    preclinical_rows,
    col_widths=[1.2, 2.8, 4.5, 2, 4, 5, 2, 1.8], font_size=8.5)

add_para(doc, '')
add_para(doc, '临床前小计', bold=True, size=11)
add_para(doc, '共 5 项；并行推进合计 6–9 个月；总费用估算 80–160 万元；均可外包给国家医疗器械检验所或 CNAS 认可第三方机构。',
         size=10, color=(89, 89, 89))

doc.add_page_break()

# ==== 二、临床研究清单 ====
add_heading(doc, '二、临床研究清单（Clinical Studies）——共 2 项核心 + 9 项支撑文献', level=1)

add_heading(doc, '2.1 核心临床试验（2 项）', level=2)

clinical_core_rows = [
    ['C-01', 'Empower（关键 pivotal RCT）',
     'IDE G210328\n\nNCT: (未公开)\n发表 Nat Med 2025',
     '前瞻、多中心、双盲、随机、伪对照 2 臂 RCT\n+ 10 周开放标签期扩展',
     '174 例（ITT）\n173 例（mITT）\nUK 115 + US 59',
     '2 家中心\nKing\'s College Hospital（英国）\nUTHealth Houston（美国）',
     '18+ MDD 当前发作期\nHDRS-17 ≥ 16\n排除 TRD\n共 15+35 条入排',
     'Active: 2.0 mA × 30 min × 36 次\n前 3 周 5 次/周\n后 7 周 3 次/周\n\nSham: ramp 1 mA/30s→0 mA/15s',
     '主要：Wk10 HDRS-17 均分改变\n次要：9 项（HDRS-17/MADRS/MADRS-s 应答/缓解 + EQ-5D-3L）',
     'MMRM 混合模型\nFCS 多重插补 20 数据集\nRubin\'s Rule 合并\nHochberg 多重性\nα=0.025 单侧',
     '主要终点显著（Δ=-2.3，p=0.012）\n应答率 54.4% vs 26.9%\n缓解率 44.9% vs 21.8%\n无器械相关 SAE',
     '14 月执行\n（22-06 → 23-08）'],

    ['C-02', 'PSYLECT（补充临床）',
     'NCT04889976\n发表 JAMA Psychiatry 2024',
     '单中心 3 臂 RCT\n双盲 sham 对照\n\n① Active tDCS + iBT\n② Active tDCS + sham iBT\n③ 双 sham',
     '210 例入组\n199 例完成',
     '1 家中心\n巴西圣保罗大学附属医院',
     '18+ MDD\nHDRS-17 ≥ 17\n（更严格阈值）',
     'Active: 2 mA × 30 min × 15 个连续工作日 + 之后 3 周每周 2 次\nSham: 1 mA × 90s',
     '主要：Wk6 HDRS-17 分数改变',
     '线性混合效应模型',
     '**阴性**\n未证明优于 sham\n所有 3 组间效应量差异不显著\n1 例非致死自杀企图（tDCS 组）',
     '未披露'],
]
add_table(doc,
    ['编号', '试验名称', '注册号/发表', '设计', '样本量', '中心', '入排要点',
     '刺激方案', '终点', '统计方法', '结果', '执行时长'],
    clinical_core_rows,
    col_widths=[1.0, 2.2, 2.2, 2.5, 1.8, 2.5, 2.3, 2.5, 2.5, 2.5, 2.5, 1.5],
    font_size=8)

add_heading(doc, '2.2 支撑文献（Flow FL-100 或同类 tDCS 独立研究 · 9 项）', level=2)

lit_rows = [
    ['L-01', 'Aparicio LVM et al. 2024',
     'PSYLECT 维持期结果', '开放标签维持期', '同 PSYLECT'],
    ['L-02', 'Borrione L et al. 2021, J Affect Disord',
     '家用 tDCS + App 心理干预治疗 MDD 的病例系列', '病例系列', '早期可行性验证'],
    ['L-03', 'Borrione L et al. 2024, JAMA Psychiatry',
     'PSYLECT 主论文（即 C-02）', 'RCT', '同 C-02'],
    ['L-04', 'Woodham RD et al. 2022, J Psychiatr Res',
     '辅助家用 tDCS 治疗 MDD 的开放标签可行性研究', '单臂开放标签 + 长期随访', '实时远程督导可行'],
    ['L-05', 'Woodham RD et al. 2025, J Psychiatr Res',
     '家用 tDCS 治疗 MDD 6 月随访（Empower 长期）', 'RCT 后续随访', 'Empower 6 月长期结果'],
    ['L-06', 'Sobral M et al. 2022, Front Psychiatry',
     '家用 tDCS 双阳性治疗抑郁焦虑症状的病例系列', '病例系列', '共病抑郁焦虑'],
    ['L-07', 'Rezaei H et al. 2025, BMC Psychiatry',
     '家用 tDCS 治疗双相抑郁的可接受性主题分析', '定性研究', '受试者体验研究'],
    ['L-08', 'Baukaite E et al. 2024, Open J Depression',
     'Flow tDCS 用于产褥期抑郁与丧偶服务对象的解释性现象学研究', '定性研究', '特殊人群探索'],
    ['L-09', 'Ghazi-Noori AR et al. 2024, Int J Bipolar Disord',
     '家用 tDCS 治疗双相抑郁的开放标签研究', '单臂开放标签', '双相扩展'],
]
add_table(doc,
    ['编号', '文献引用', '内容', '研究类型', '注册价值'],
    lit_rows,
    col_widths=[1.0, 4.5, 8, 4, 6], font_size=9)

add_heading(doc, '2.3 引用的方法学文献（3 项，用于方案设计参考）', level=2)

method_rows = [
    ['M-01', 'Matsumoto & Ugawa 2016, Clin Neurophysiol Pract',
     'tDCS 和 tACS 不良反应综述——AE 收集清单来源'],
    ['M-02', 'Brunoni AR et al. 2011, Int J Neuropsychopharmacol',
     'tDCS 不良反应报告与评估系统综述——AEQ 量表来源'],
    ['M-03', 'Szigeti B et al. 2023, Sci Rep',
     'placebo 组 vs placebo 对照的差异——CGR 设盲调整方法来源'],
    ['M-04', 'Brunoni AR et al. 2017, NEJM (ELECT-TDCS)',
     '经典 tDCS vs 依他普仑 抗抑郁头对头 RCT——中国方案可作对照参考'],
    ['M-05', 'Herdman M et al. 2011, Qual Life Res',
     'EQ-5D-5L 开发与验证——生活质量量表选择依据'],
]
add_table(doc,
    ['编号', '文献引用', '内容/用途'],
    method_rows,
    col_widths=[1.5, 6, 16.5], font_size=9)

doc.add_page_break()

# ==== 三、Starstim 中国工作分解（本报告新增）====
add_heading(doc, '三、Starstim 中国注册工作分解表（本报告综合建议）', level=1)

work_rows = [
    # 阶段, 编号, 工作项, 内容, 依据/参考, 预计时间, 预计费用, 负责方
    ['前期', 'W-01', 'CDE Pre-submission 沟通',
     '分类界定（三类 09-03）+ 适应症锁定（推荐 MDD 辅助治疗）+ 临床方案框架确认',
     '国家药监局政务服务门户',
     '2–3 月',
     '5–15 万',
     '注册代理 + 内部注册组'],

    ['前期', 'W-02', '技术转让协议（与 Neuroelectrics）',
     '设计文件、图纸、软件源代码、检验规范、生产工艺、专利许可、临床数据可引用性',
     'Emergo by UL 中国注册指南',
     '3–6 月（与 CDE 沟通并行）',
     '30–80 万（法务）',
     '法务 + Neuroelectrics 谈判'],

    ['临床前', 'W-03', '产品技术要求编制',
     '覆盖 5 项非临床检验维度',
     '对标 Flow FL-100 SSED §X',
     '2 月',
     '5–10 万',
     '注册代理 + 内部研发'],

    ['临床前', 'W-04', '生物相容性（P-01）',
     '细胞毒性 + 皮肤致敏 + 皮内反应',
     'GB/T 16886.5/10/23',
     '2–4 月',
     '15–30 万',
     '国家医疗器械检验所或 CNAS 第三方'],

    ['临床前', 'W-05', 'EMC + 电气安全（P-02）',
     'GB 9706.1 + YY 9706.102 + YY 9706.111 + YY 9706.210',
     '同上标准',
     '4–6 月',
     '25–45 万',
     '同上'],

    ['临床前', 'W-06', '软件 + 网络安全（P-03）',
     '软件生存周期文档 + 网络安全测试',
     'YY/T 0664 + YY/T 1843',
     '3–6 月（与研发并行）',
     '20–40 万',
     '软件测评机构（如中软测评、赛宝）'],

    ['临床前', 'W-07', '有效期验证（P-04）',
     '加速老化 + 实时老化',
     'YY/T 0681 参考',
     '6–12 月',
     '10–20 万',
     '内部实验室或第三方'],

    ['临床前', 'W-08', '出厂性能验证（P-05）',
     '每台电流精度 100% 检验',
     '产品技术要求',
     '3–6 月（含产线搭建）',
     '10–25 万',
     '内部生产'],

    ['临床', 'W-09', '临床方案 + SAP 编写',
     '基于 Empower 一比一映射，本土化改造',
     'FDA SSED §XI + 依瑞德案例',
     '3 月',
     '15–30 万',
     'CRO + 精神科专家'],

    ['临床', 'W-10', '伦理审查（多中心）',
     '3 家三甲医院精神科伦理会',
     'GCP',
     '3 月（并行）',
     '9–15 万（3 中心）',
     '临床中心 + CRO'],

    ['临床', 'W-11', '主临床试验（C-01 对应）',
     '3 中心 × 150 例 × 10 周盲态 + 10 周开放；主要终点 HAMD-17 均分改变',
     '基于 Empower 设计',
     '15–18 月（含入组+治疗+随访）',
     '400–600 万',
     'CRO + 3 家中心'],

    ['临床', 'W-12', '统计分析与总结报告',
     'MMRM + FCS 插补 + Hochberg 多重性调整；CGR 设盲敏感性分析',
     '基于 Empower 统计方法',
     '3–4 月',
     '20–40 万',
     '独立统计单位'],

    ['临床', 'W-13', '临床评价报告（CER）',
     '整合境外数据 + 中国 RCT 数据 + 与 TMS/ECT 对比',
     'NMPA 2021 年第 73 号通告',
     '2 月',
     '8–15 万',
     '注册代理'],

    ['注册', 'W-14', '注册资料整理 + 递交',
     '含技术要求、检验报告、CER、说明书、风险管理、体系文件等',
     'NMPA 医疗器械注册资料要求',
     '2 月',
     '5–10 万',
     '注册代理'],

    ['注册', 'W-15', '技术审评 + 发补 + 体系核查',
     '1–2 轮发补应对；境内生产企业现场体系核查',
     'CDE 技术审评流程',
     '12–18 月',
     '30–50 万（发补应对）',
     '注册代理 + 现场配合'],

    ['注册', 'W-16', '注册证发放',
     '国械注准 20xxxxxx',
     'NMPA',
     '1–2 月',
     '15.36 万（三类境内注册官费）',
     'NMPA'],
]
add_table(doc,
    ['阶段', '编号', '工作项', '内容', '依据/参考', '预计时间', '预计费用', '负责方'],
    work_rows,
    col_widths=[1.5, 1.2, 3.5, 6, 3.5, 3, 3, 3.5], font_size=8.5)

add_para(doc, '')
add_para(doc, '中国注册总计', bold=True, size=11, color=(192, 0, 0))
add_para(doc,
    '总工作项 16 项；关键路径 30–48 月（临床 W-11 主导）；总费用估算 640–1090 万元（不含"国产化落地"隐性成本）；'
    '合计与《引进国产化注册可行性评估报告》基线预算一致。',
    size=10, color=(89, 89, 89))

# ==== 四、一句话对比 ====
add_heading(doc, '四、Flow FL-100 vs Starstim 中国方案 一句话对比', level=1)

compare_rows = [
    ['非临床研究项数', '5 项', '5 项（一一对应）', '完全对齐'],
    ['关键 RCT 中心数', '2 家（英美各 1）', '3 家（推荐三甲精神科）', '中国 CDE 惯例更严'],
    ['关键 RCT 样本量', '174 例', '150 例', '中国方案已足够'],
    ['盲态期长度', '10 周', '10 周', '完全一致'],
    ['开放标签期', '10 周', '10 周', '完全一致'],
    ['主要终点', 'HDRS-17 均分改变', 'HAMD-17 均分改变（等价）', '术语本土化'],
    ['统计方法', 'MMRM + FCS + Hochberg', '同上', '完全一致'],
    ['执行模式', '全远程家用', '半远程（保留 4 次现场访视）', '中国本土化'],
    ['预计执行时长', '14 月', '15–18 月', '略长'],
    ['关键风险', '设盲失败', '设盲失败 + CDE 发补', '预设 CGR 敏感性分析应对'],
    ['预计总费用', '未披露', '400–600 万（主临床）', '一次性投入'],
]
add_table(doc,
    ['维度', 'Flow FL-100（实际做法）', 'Starstim 中国建议方案', '备注'],
    compare_rows,
    col_widths=[4, 6, 6, 6], font_size=10)

# 结尾
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('—— 清单完 · 共 3 类研究 (5+2+9) + 16 项工作分解 ——')
set_run_font(r, size=11, color=(89, 89, 89))

# 保存
out = '/home/user/sugz/scratchpad/Flow_FL100_研究清单表_工作分解.docx'
doc.save(out)
print(f'Saved: {out}')
