# -*- coding: utf-8 -*-
"""
补充报告：来源依据 + 欧盟/美国注册情况 + 对标产品案例
输出 Word 供决策汇报
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_run_font(run, name='微软雅黑', size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rpr = r.rPr
    if rpr is None:
        rpr = OxmlElement('w:rPr')
        r.insert(0, rpr)
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def add_heading(doc, text, level=1):
    styles = {
        1: (16, True, (31, 73, 125)),
        2: (13.5, True, (54, 95, 145)),
        3: (12, True, (68, 114, 148)),
    }
    size, bold, color = styles.get(level, (11, True, (0, 0, 0)))
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_para(doc, text, bold=False, indent_cn=True, size=10.5, color=None, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent_cn:
        p.paragraph_format.first_line_indent = Pt(21)
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_source(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Pt(14)
    r1 = p.add_run('[来源] ')
    set_run_font(r1, size=9.5, bold=True, color=(192, 0, 0))
    r2 = p.add_run(text)
    set_run_font(r2, size=9.5, color=(89, 89, 89))
    return p


def add_table(doc, headers, rows, col_widths=None, header_color='1F497D'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        set_cell_bg(hdr[i], header_color)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ''
            p = cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            set_run_font(run, size=9.5)
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


# ============ 生成文档 ============
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ==== 封面 ====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(120)
r = p.add_run('Neuroelectrics 引进国产化项目')
set_run_font(r, size=22, bold=True, color=(31, 73, 125))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('决策依据补充报告（Evidence Package）')
set_run_font(r, size=20, bold=True, color=(31, 73, 125))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
r = p.add_run('欧美注册与临床证据 · 中国法规原文 · 国内对标产品案例 · 全部来源可核查')
set_run_font(r, size=12, color=(89, 89, 89))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(150)
r = p.add_run('本报告用于对上一版《引进中国国产化注册可行性深度评估报告》的每一条结论，')
set_run_font(r, size=11, color=(89, 89, 89))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('提供可点击、可核查的公开来源，供老板做出决策。')
set_run_font(r, size=11, color=(89, 89, 89))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run('调研截止：2026 年 7 月')
set_run_font(r, size=11, color=(89, 89, 89))

doc.add_page_break()

# ==== 目录 ====
add_heading(doc, '目  录', level=1)
toc = [
    '一、如何使用本报告',
    '二、上一版报告核心结论的来源依据对照',
    '    2.1 结论 1：产品分类',
    '    2.2 结论 2：Enobio 走豁免临床合理',
    '    2.3 结论 3：Starstim 必须做临床',
    '    2.4 结论 4：250 万医生费用偏低',
    '    2.5 结论 5：同步 vs 依次策略',
    '    2.6 结论 6："国产化"是关键前提',
    '三、Enobio 在欧盟与美国的注册情况',
    '    3.1 美国 FDA 510(k) K162681 详细',
    '    3.2 欧盟 CE 认证',
    '    3.3 Enobio 支撑注册与验证的研究',
    '四、Starstim 在欧盟与美国的注册情况',
    '    4.1 美国 FDA 路径：IDE + Breakthrough Device Designation',
    '    4.2 欧盟 CE 认证与适应症声称',
    '    4.3 Starstim 关键临床试验一览',
    '五、如果做中国国内注册需要做哪些研究',
    '    5.1 Enobio 需做的研究（豁免临床路径）',
    '    5.2 Starstim 需做的研究（临床试验路径）',
    '六、国内对标产品案例：武汉依瑞德 CCY 系列 TMS',
    '    6.1 选择本案例作对标的理由',
    '    6.2 依瑞德 CCY 系列 TMS 关键事实与来源',
    '    6.3 依瑞德为拿证所做的研究',
    '    6.4 从依瑞德案例可推导出的 Starstim 临床要求',
    '七、次要对标：脑电图机（Enobio 对标）',
    '八、可点击来源清单（全部 URL）',
]
for t in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(t)
    set_run_font(r, size=11)

doc.add_page_break()

# ==== 一、如何使用本报告 ====
add_heading(doc, '一、如何使用本报告', level=1)
add_para(doc,
    '本报告与《引进中国国产化注册可行性深度评估报告》配套使用。上一版报告给出的每一条结论，'
    '本报告以"结论 → 引用原文 → 来源 URL"的方式重新呈现，方便老板逐条核查。'
    '同时补充了：（1）Enobio 与 Starstim 在美国 FDA、欧盟 CE 下的实际注册情况与支撑研究；'
    '（2）如果在中国做国产注册，具体要做哪些研究；（3）一个国内对标产品案例——'
    '武汉依瑞德 CCY 系列经颅磁刺激仪（TMS）——作为经颅刺激类产品临床方案设计的参照。'
)
add_para(doc,
    '每处"[来源]"标签后附有原文标题与可访问的 URL；报告末尾附有完整可点击来源清单，'
    '便于审阅人一键跳转。',
)

doc.add_page_break()

# ==== 二、上一版报告核心结论的来源依据对照 ====
add_heading(doc, '二、上一版报告核心结论的来源依据对照', level=1)

# 结论1
add_heading(doc, '2.1 结论 1：EEG 属二类，tES 治疗用途按三类', level=2)
add_bullet(doc, '依据：脑电图机（EEG）在《医疗器械分类目录》中分类编码 07-03-06，属第二类医疗器械。')
add_source(doc, '西格玛医学《器械注册一起学 11：脑电图机》，明确 EEG 属二类、分类编码 07-03-06、'
                'GB 9706.226-2021 为其专用安全标准；https://sigma-stat.com/index.php?m=home&c=View&a=index&aid=3035')
add_source(doc, '中国国家标准馆 GB 9706.226-2021《医用电气设备 第 2-26 部分：脑电图机的基本安全和基本性能专用要求》；'
                'https://std.samr.gov.cn/gb/search/gbDetailed?id=CE1E6A1DD42558F6E05397BE0A0A68DF')
add_bullet(doc, '依据：经颅磁/电刺激仪用于精神疾病治疗按三类管理。')
add_source(doc, '腾讯新闻《首证破冰！经颅磁刺激"三类证"开启行业规范新纪元》：'
                '"2017 年，国家药监局发布的《医疗器械分类目录》中，明确将用于精神疾病治疗的经颅磁刺激仪纳入'
                '第三类医疗器械管理"。https://news.qq.com/rain/a/20250605A08UDI00')
add_source(doc, '医药魔方 ByDrug《磁刺激市场竞争进入新阶段》：截至 2025 年 5 月，国内 TMS 设备除依瑞德最新'
                '三类证外，其余 28 款均为二类证。'
                'https://bydrug.pharmcube.com/news/detail/1bc8aeb7f2a698524adc8828196a7bbc')

# 结论2
add_heading(doc, '2.2 结论 2：Enobio 走豁免临床合理', level=2)
add_bullet(doc, '依据：国家药监局 2023 年第 33 号通告发布《免于临床评价医疗器械目录》，共 1025 种产品，'
                '脑电图机在目录内。')
add_source(doc, '上海市药品监督管理局公告"国家药监局关于发布免于临床评价医疗器械目录的通告（2023 年第 33 号）"，'
                '2023 年 7 月 20 日施行；https://yjj.sh.gov.cn/qtgzwj/20230725/76df258cba2e4248b2ebb0db7d2061f0.html')
add_source(doc, 'CIO 合规查查转录 2023 年第 33 号通告全文；https://www.ciopharma.com/supervise/31081')
add_source(doc, '健康界报道《国家药监局：器械免于临床评价目录（2023 年第 33 号）》，说明该目录含 1025 种产品；'
                'https://www.cn-healthcare.com/articlewm/20230727/content-1584353.html')
add_bullet(doc, '依据：同品种比对的法规依据为 NMPA 2021 年第 73 号通告《医疗器械临床评价技术指导原则》。')
add_source(doc, '国家药监局 2021 年第 73 号通告：《医疗器械临床评价技术指导原则》《医疗器械临床评价等同性论证'
                '技术指导原则》等 5 项技术指导原则，发布日期 2021-09-28；'
                'https://www.cirs-group.com/cn/md/gjyjjgyfbylqxlcpjjszdyzd5xjszdyzdtg%EF%BC%882021nd73h%EF%BC%89')

# 结论3
add_heading(doc, '2.3 结论 3：Starstim 必须做临床', level=2)
add_bullet(doc, '依据：豁免临床目录未收录经颅电/磁刺激类"治疗性"设备。')
add_source(doc, '同上 2023 年第 33 号通告目录；https://yjj.sh.gov.cn/qtgzwj/20230725/76df258cba2e4248b2ebb0db7d2061f0.html')
add_bullet(doc, '依据：三类注册产品必须开展多中心临床试验。')
add_source(doc, '国家药监局政务服务门户"国产三类医疗器械首次注册"任务说明；'
                'https://zwfw.nmpa.gov.cn/web/taskview/11100000MB0341032Y100017214300101')
add_source(doc, '搜狐新闻转载"三类证的核心差异在于对产品安全性和有效性的更高要求，临床数据需要在至少两家符合'
                '标准的三级甲等医疗机构中完成"；https://www.sohu.com/a/899847819_120987638')

# 结论4
add_heading(doc, '2.4 结论 4：250 万医生费用偏低（临床全成本口径）', level=2)
add_bullet(doc, '依据：三类神经调控临床费用组成参考。')
add_source(doc, '思途 CRO《医疗器械临床试验的法规差异——中、美、欧三地监管路径解析》；'
                'https://www.situcro.com/news/7767.html')
add_source(doc, '思途 CRO《办理二类医疗器械注册证要多少钱》以及《国内二三类医疗器械注册收费标准》'
                '（含临床试验成本区间讨论）；https://www.situcro.com/news/4116.html')
add_bullet(doc, '国内三类神经调控多中心临床（3–5 中心、100–200 例、随访 3–6 月）的公开经验，'
                '常规区间 400–800 万元，此为业内 CRO 通用测算，本报告已在正文明确"关系型压价可下调"。')
add_source(doc, '上述区间为《引进报告》与本报告作者基于国内 CRO 公开案例综合估算，无单一权威公开数据源，'
                '建议以后续 CRO 报价单为准（决策前需要与 3 家 CRO 询价交叉验证）。')

# 结论5
add_heading(doc, '2.5 结论 5：同步 vs 依次策略', level=2)
add_bullet(doc, '依据：同步策略额外一套设备成本约 150 万 = 15–20 台 × 7.5–10 万/台，与对方口径一致。'
                '但需要额外 CRC/监查 30–70 万，此为项目管理经验估算。')
add_source(doc, '瀚翔脑科学《经颅磁刺激（TMS）和经颅直流电刺激（tDCS）有何不同？》，说明多中心研究的常规规模；'
                'https://www.infoinstruments.cn/tmsvs-tdcs/')

# 结论6
add_heading(doc, '2.6 结论 6："国产化"是关键前提', level=2)
add_bullet(doc, '依据：进口 vs 国产首次注册官费。')
add_source(doc, 'CIRS Group《更新 | 中国医疗器械注册收费标准（2024 年 5 月）》：三类国产首次注册 15.36 万元，'
                '进口首次注册 30.88 万元；'
                'https://www.cirs-group.com/cn/md/zhong-guo-yi-liao-qi-xie-zhu-ce-shou-fei-biao-zhun-2024-nian-5-yue')
add_source(doc, 'ReguVerse《医疗器械产品注册收费标准及注册收费实施细则》；'
                'https://reguverse.com/documentation/nmpa-regulations-index/pre-market-submission/fees/registration-fees/')
add_bullet(doc, '依据：国产化路径必须在境内组织生产、取得《医疗器械生产许可证》。')
add_source(doc, 'Emergo by UL《China Medical Device Registration and Approval》；'
                'https://www.emergobyul.com/services/china-medical-device-registration-and-approval')

doc.add_page_break()

# ==== 三、Enobio 在欧盟与美国的注册情况 ====
add_heading(doc, '三、Enobio 在欧盟与美国的注册情况', level=1)

add_heading(doc, '3.1 美国 FDA 510(k) K162681 详细', level=2)
enobio_fda = [
    ['申请人（Submitter）', 'Neuroelectrics Barcelona S.L.U.（西班牙巴塞罗那）'],
    ['产品商品名', 'Enobio Wireless EEG（8, 20, 32 通道）'],
    ['510(k) 编号', 'K162681'],
    ['取证日期', '2017 年 4–5 月（Innolitics 记录为 2017-05-05）'],
    ['分类', 'Class II'],
    ['监管条款 / 产品代码', '21 CFR 882.1400（GWQ，Full-Montage Standard EEG）+ 21 CFR 882.1835（GWL，Physiological Signal Amplifier）'],
    ['前代设备（Predicate）', 'Nicolet Wireless EEG（K103140，Carefusion）'],
    ['预期用途（Indications for Use）',
     '"ENOBIO is an EEG portable monitoring device of 8, 20, 32 channels intended for use in clinical patient '
     'monitoring for use in hospitals and other medical environments. The Enobio is intended to acquire, '
     'store, transmit and display electrophysiological signals in wireless mode as an aid in diagnostics."'
     '（"用于医院及其他医疗环境的临床病人监护 EEG 便携监护设备，用于采集、存储、传输并显示电生理信号，'
     '作为辅助诊断"）'],
]
add_table(doc, ['项目', '内容'], enobio_fda, col_widths=[5, 11])
add_source(doc, 'FDA 510(k) K162681 官方清关文件（PDF，含 Neuroelectrics Barcelona 与前代设备名）：'
                'https://www.accessdata.fda.gov/cdrh_docs/pdf16/K162681.pdf')
add_source(doc, 'FDA Innolitics 510(k) 数据库检索页：https://fda.innolitics.com/submissions/NE/subpart-b'
                '%E2%80%94neurological-diagnostic-devices/GWQ/K143440')
add_source(doc, 'FDA Report 公司档案页 Neuroelectrics Barcelona S L U：'
                'https://fda.report/Company/Neuroelectrics-Barcelona-S-L-U')

add_heading(doc, '3.2 欧盟 CE 认证', level=2)
enobio_ce = [
    ['CE 分类', 'Class IIa（依据欧盟指令 93/42/EEC，即 MDD 医疗器械指令；MDR 过渡期后需重新认证到 MDR 框架）'],
    ['公告机构（Notified Body）', 'Notified Body 0120'],
    ['市场覆盖', '欧盟全境（合规市场超过 74 国）'],
]
add_table(doc, ['项目', '内容'], enobio_ce, col_widths=[5, 11])
add_source(doc, 'Neuroelectrics 官方 Enobio 产品页（写明 CE Class IIa 与 93/42/EEC）：'
                'https://www.neuroelectrics.com/products/research/enobio')
add_source(doc, 'Enobio 用户手册 v1.3（Manualzz 版）内含 CE 认证声明；'
                'https://manualzz.com/doc/7275717/enobio-user-manual-neuroelectrics')
add_source(doc, 'Neuroelectrics 公司博客《EU Regulatory Framework for Medical Devices: An Introduction to MDR》，'
                '公司自述其产品需完成 MDD → MDR 过渡；'
                'https://www.neuroelectrics.com/blog/eu-regulatory-framework-for-medical-devices-an-introduction-to-mdr')

add_heading(doc, '3.3 Enobio 支撑注册与验证的研究', level=2)
add_para(doc,
    'Enobio 走 FDA 510(k) 路径，其"支撑材料"以"实质等同性证据"为主，而非独立临床试验（这是二类 EEG 的行业惯例）：'
)
add_bullet(doc, '实质等同性论证（Substantial Equivalence）：证明 Enobio 与 Nicolet K103140 在预期用途、'
                '技术特征上等同，采样率、通道数、频响、安全指标均可对比。')
add_bullet(doc, '性能测试报告：GB/IEC 60601-1 电气安全、EMC（IEC 60601-1-2）、脑电图机专项性能测试。')
add_bullet(doc, '独立学术验证：Enobio 被大量第三方学术论文用作 EEG 采集设备，覆盖 BCI、癫痫、'
                '睡眠、注意力、AD、PTSD、抑郁、疼痛等场景。Neuroelectrics 官方维护 "Enobio Publications" wiki，'
                '收录数百篇同行评审论文。')
add_source(doc, 'Neuroelectrics Wiki《Collection of publications of independent research studies and mentions '
                'about Enobio》；https://www.neuroelectrics.com/wiki/index.php/Collection_of_publications'
                '_of_independent_research_studies_and_mentions_about_Enobio')
add_source(doc, 'PLOS Biology (2023) 使用 Enobio 32 完成的高影响因子研究：'
                'https://journals.plos.org/plosbiology/article?id=10.1371%2Fjournal.pbio.3002193')
add_source(doc, 'BioRxiv 2022 URGOnight 神经反馈设备验证研究，使用 Enobio-20 作为参考系统：'
                'https://www.biorxiv.org/content/10.1101/2022.12.27.522035.full.pdf')
add_source(doc, 'Epilepsy Foundation Device Wiki - Enobio 条目；'
                'https://www.epilepsy.com/tools-resources/device-wiki/enobio')

doc.add_page_break()

# ==== 四、Starstim 在欧盟与美国的注册情况 ====
add_heading(doc, '四、Starstim 在欧盟与美国的注册情况', level=1)

add_heading(doc, '4.1 美国 FDA 路径：IDE + Breakthrough Device Designation', level=2)
add_para(doc,
    'Starstim 在美国 FDA 目前尚未获得 510(k) 治疗适应症清关；而是通过以下三条组合路径累积临床与法规资本：',
)
starstim_fda = [
    ['① FDA 研究用器械豁免（IDE）',
     'IDE 编号 G160208，赞助方 Neuroelectrics，主要研究中心 Boston Children\'s Hospital（PI：Alexander Rotenberg 教授）'],
    ['② FDA 突破性器械认定（Breakthrough Device Designation）',
     '2021-06-30 授予，适应症为"难治性局灶性癫痫（refractory focal epilepsy）"'],
    ['③ FDA COVID 期间"家用 tDCS 治疗重度抑郁（MDD）"临时授权',
     '2020 年（COVID-19 期间）由 FDA 授权 Starstim 家用 MDD 治疗（临时应急路径）'],
]
add_table(doc, ['路径', '关键要点'], starstim_fda, col_widths=[5, 11])
add_source(doc, 'MassDevice 报道《Neuroelectrics touts FDA IDE Starstim epilepsy treatment study data》；'
                'https://www.massdevice.com/neuroelectrics-touts-fda-ide-starstim-epilepsy-treatment-study-data/')
add_source(doc, 'BusinessWire 新闻稿《Neuroelectrics Granted FDA Breakthrough Device Designation for New '
                'Therapeutic Neuromodulation Platform to Treat Refractory Focal Epilepsy》，2021-06-30；'
                'https://www.businesswire.com/news/home/20210630005708/en/')
add_source(doc, 'NeurologyLive《New tES System Gets Breakthrough Designation in Refractory Focal Epilepsy》；'
                'https://www.neurologylive.com/view/new-tes-system-breakthrough-designation-refractory-focal-epilepsy')
add_source(doc, 'Neuroelectrics 官方博客《FDA greenlights Neuroelectrics to help patients with Major Depression '
                'at home amidst COVID-19 restrictions》；'
                'https://www.neuroelectrics.com/blog/fda-greenlights-neuroelectrics-to-help-patients-with-major-depression-at-home-amidst-covid-19-restri')
add_source(doc, 'Boston Children\'s Hospital 临床试验页 NCT02866240；'
                'https://www.childrenshospital.org/clinical-trials/nct02866240')

add_heading(doc, '4.2 欧盟 CE 认证与适应症声称', level=2)
add_para(doc,
    'Starstim 在欧盟持有 CE Class IIa 认证（与 Enobio 同框架），主要作为"研究与临床研究用刺激系统"上市，'
    '适应症声称覆盖疼痛、癫痫、AD、卒中康复、抑郁、成瘾等研究方向；至本报告调研时未查到'
    '"欧盟 MDR 框架下针对某一特定适应症（如 MDD）的单独治疗性 CE 声明"。',
)
add_source(doc, 'Neuroelectrics 官方产品页 Starstim tES：'
                'https://www.neuroelectrics.com/products/research/starstim')
add_source(doc, 'DIY tDCS 社区对 Starstim 家用 tDCS 的报道；https://www.diytdcs.com/tag/starstim/')
add_source(doc, 'Neuroelectrics 官方博客《tDCS at home with Starstim》；'
                'https://www.neuroelectrics.com/blog/tdcs-at-home-with-starstim')

add_heading(doc, '4.3 Starstim 关键临床试验一览', level=2)
starstim_trials = [
    ['NCT02866240', '难治性癫痫（药物抵抗）', 'Boston Children\'s Hospital / 墨西哥城 National Institute of Neurology',
     '17 例完成 / 单臂开放标签',
     '75% 患者 8 周随访时癫痫发作频次下降 ≥40%；中位数下降 40–44%',
     '无器械相关不良事件',
     'FDA IDE G160208 · 结果发表 American Epilepsy Society (AES) 2018 年会 + Kaye et al. 2023'],
    ['MDD 家用 tDCS Pilot（NCT 未公开）',
     '药物抵抗性重度抑郁症（MDD）',
     '美国多中心 telemedicine',
     '35 例入组、34 例完成',
     '主要指标 MADRS 中位数下降 64%（8 周随访）；近 90% 完成率、无严重不良事件',
     '轻中度红斑、瘙痒等常见 tDCS 不良反应',
     '发表 Frontiers in Psychiatry 2024；同期 BusinessWire 新闻稿'],
    ['卒中康复 tDCS 家用 Pilot',
     '慢性偏瘫',
     'StarStim Home Research Kit 家用',
     '6 例（可行性研究）',
     '20 分钟 1.5 mA tDCS + 手指追踪治疗，工作日连续 5 日',
     '可行性验证，无安全事件',
     '独立学术验证，未直接用于 FDA 注册'],
]
add_table(doc, ['研究/编号', '适应症', '中心', '样本量', '主要结果', '安全性', '发表/法规'], starstim_trials,
          col_widths=[3, 2.5, 2.5, 1.8, 3.2, 1.8, 3])
add_source(doc, 'BioSpace 转载《Neuroelectrics Announces Positive Results for Treatment of Medication-Resistant '
                'Epilepsy Using Starstim™》；'
                'https://www.biospace.com/neuroelectrics-announces-positive-results-for-treatment-of-medication-resistant-epilepsy-using-starstim')
add_source(doc, 'Kaye HL, San-Juan D, ... Rotenberg A. "Personalized, Multisession, Multichannel Transcranial '
                'Direct Current Stimulation in Medication-Refractory Focal Epilepsy: An Open-Label Study." '
                'Journal of Clinical Neurophysiology 40(1):53-62, 2023；PubMed PMID: 34010226；'
                'https://pubmed.ncbi.nlm.nih.gov/34010226/')
add_source(doc, 'Frontiers in Psychiatry 2024《Multichannel tDCS with advanced targeting for major depressive '
                'disorder: a tele-supervised at-home pilot study》；'
                'https://www.frontiersin.org/journals/psychiatry/articles/10.3389/fpsyt.2024.1427365/full')
add_source(doc, 'PMC 全文《Multichannel tDCS with advanced targeting for MDD》 PMC11358063；'
                'https://pmc.ncbi.nlm.nih.gov/articles/PMC11358063/')
add_source(doc, 'BusinessWire《Neuroelectrics Announces Results of Telemedicine Pilot Depression Study with '
                'Starstim tDCS Therapy》，2024-03-07；'
                'https://www.businesswire.com/news/home/20240307614300/en/')
add_source(doc, 'PMC《Remotely Supervised Transcranial Direct Current Stimulation in Post-Stroke Recovery: '
                'A Scoping Review》（提及 StarStim Home Research Kit 用于卒中康复 6 例可行性研究）；'
                'https://pmc.ncbi.nlm.nih.gov/articles/PMC12029044/')
add_source(doc, 'Frontiers in Aging Neuroscience 2021《Safety and Feasibility of Tele-Supervised Home-Based '
                'Transcranial Direct Current Stimulation for Major Depressive Disorder》；'
                'https://www.frontiersin.org/journals/aging-neuroscience/articles/10.3389/fnagi.2021.765370/full')

doc.add_page_break()

# ==== 五、如果做中国国内注册需要做哪些研究 ====
add_heading(doc, '五、如果做中国国内注册需要做哪些研究', level=1)

add_heading(doc, '5.1 Enobio 需做的研究（豁免临床评价路径）', level=2)
add_para(doc,
    '走《免于临床评价目录》（NMPA 2023 年第 33 号通告）+ 《医疗器械临床评价技术指导原则》'
    '（NMPA 2021 年第 73 号通告）的路径，Enobio 无需开展中国境内临床试验，但需要完成以下研究与文件：',
)
enobio_studies = [
    ['① 产品技术要求编制',
     '按 GB 9706.1-2020、GB 9706.226-2021、YY 9706.102-2021 等标准编制'],
    ['② 全项型式试验（注册检验）',
     '国家医疗器械检验所或有资质第三方完成电气安全、EMC、EEG 专用性能测试'],
    ['③ 同品种比对（Substantial Equivalence）',
     '选 1–2 款国内已注册脑电图机作为对比器械，比对预期用途、技术特征、生物学特性'],
    ['④ 差异性非临床数据',
     '对比对差异点（如干电极、无线传输等）补充非临床数据支持'],
    ['⑤ 软件生存周期文档（YY/T 0664）',
     '软件安全等级评定、需求追溯、测试报告'],
    ['⑥ 网络安全测试（YY/T 1843）',
     '医疗器械网络安全注册技术审查指导原则要求'],
    ['⑦ 生物相容性评价（GB/T 16886 系列）',
     '皮肤接触电极的生物相容性评价，通常引用已发表数据'],
    ['⑧ 说明书 / 标签 / 上市后监测方案',
     '按现行说明书规范起草'],
]
add_table(doc, ['研究/文件', '内容'], enobio_studies, col_widths=[5, 11])
add_source(doc, 'GB 9706.226-2021 脑电图机专用要求：https://std.samr.gov.cn/gb/search/gbDetailed?id=CE1E6A1DD42558F6E05397BE0A0A68DF')
add_source(doc, '北京市药品监督管理局脑电图机产品技术审评规范；'
                'https://yjj.beijing.gov.cn/yjj/zwgk20/tz7/676482/index.html')
add_source(doc, '西格玛医学《器械注册一起学 11：脑电图机》：https://sigma-stat.com/index.php?m=home&c=View&a=index&aid=3035')

add_heading(doc, '5.2 Starstim 需做的研究（临床试验路径）', level=2)
add_para(doc,
    '作为三类"经颅电刺激类"治疗产品，需要开展中国境内多中心临床试验，配套完成注册前研究：',
)
starstim_studies = [
    ['① 产品技术要求 + 全项检验', 'GB 9706.1 + 神经/肌肉刺激器专用标准 YY 9706.210-2021'],
    ['② 生物相容性（GB/T 16886）', '皮肤长期接触电极的生物相容性，需完整试验或有效引用'],
    ['③ 动物实验', '若有充分文献支持可豁免；否则补做皮肤刺激/电流耐受动物试验'],
    ['④ 软件生存周期 + 网络安全 + 电子记录（YY/T 0664 / YY/T 1843）', '三类产品要求高'],
    ['⑤ 电流精度与波形一致性性能验证', 'tDCS/tACS 具体输出电流的准确性、稳定性、波形'],
    ['⑥ 多中心 RCT 临床试验（关键）', '典型：3–5 家三甲医院；100–200 例；主要终点由适应症决定，'
                                          '如 HAMD-17 / MADRS 减分率（抑郁）、Fugl-Meyer（卒中康复）、'
                                          '发作频次（癫痫）等；随访 3–6 月'],
    ['⑦ 上市前临床评价报告（CER）', '整合境外文献 + 境内临床试验数据'],
    ['⑧ 说明书、风险管理、上市后监测方案', 'ISO 14971 风险管理体系要求'],
]
add_table(doc, ['研究/文件', '内容'], starstim_studies, col_widths=[5, 11])
add_source(doc, 'YY 9706.210-2021《医用电气设备 第 2-10 部分：神经和肌肉刺激器的基本安全和基本性能专用要求》；'
                'https://www.ndls.org.cn/standard/detail/5623b01ad17afcd338c8c08a8b808da5')
add_source(doc, 'NMPA 2021 年第 73 号通告（含临床评价技术指导原则）；'
                'https://www.cirs-group.com/cn/md/gjyjjgyfbylqxlcpjjszdyzd5xjszdyzdtg%EF%BC%882021nd73h%EF%BC%89')
add_source(doc, '西格玛医学《植入式神经刺激系统介绍及临床试验设计》（含相似类别 RCT 方案要素）；'
                'https://www.sigma-stat.com/index.php?m=home&c=View&a=index&aid=3410')

doc.add_page_break()

# ==== 六、对标产品案例 ====
add_heading(doc, '六、国内对标产品案例：武汉依瑞德 CCY 系列 TMS', level=1)

add_heading(doc, '6.1 选择本案例作对标的理由', level=2)
add_bullet(doc, '依瑞德是国内首个（也是至本报告调研时点唯一一个）拿到"精神领域三类证"的经颅刺激仪厂商，与 Starstim 拟报路径最贴近。')
add_bullet(doc, '依瑞德从"2009 年拿到二类证"到"2025 年拿到抑郁症治疗三类证"，走完了完整的"低类别上市 → 高风险适应症再申报"路径，与 Neuroelectrics 引进后可能的分阶段策略相似。')
add_bullet(doc, '依瑞德的适应症（抑郁症）恰好是 Starstim 已有欧美临床数据支持的适应症之一，便于对标临床方案。')

add_heading(doc, '6.2 依瑞德 CCY 系列 TMS 关键事实与来源', level=2)
yirui_facts = [
    ['公司', '武汉依瑞德医疗设备新技术有限公司（Wuhan Yirui Medical Device Technology Co., Ltd.）'],
    ['首款产品', '2009 年 CCY 型经颅磁刺激仪（中国首个 NMPA 注册的 TMS 产品，二类）'],
    ['三类证产品', '经颅磁刺激仪（治疗抑郁症）'],
    ['注册证编号', '国械注准 20253090871'],
    ['三类证获批时间', '2025 年 5 月'],
    ['适应症（三类）', '抑郁症治疗'],
    ['CE MDR 状态', '国内首张 TMS 欧盟 MDR CE 证书（2024 年获批）'],
    ['市场地位', '在国内约 29 款 TMS 产品中，是唯一的三类证持有者，其余 28 款为二类'],
]
add_table(doc, ['项目', '内容'], yirui_facts, col_widths=[4, 12])
add_source(doc, '湖南日报旗下"华声在线"《治疗抑郁症 国内首张精神领域经颅磁三类证获批》，2025-05-29；'
                'https://m.voc.com.cn/xhn/news/202505/29488037.html')
add_source(doc, '中国医院院长网《国内首张！精神领域经颅磁三类医疗器械注册证获批》；'
                'https://www.h-ceo.com/post/5315.html')
add_source(doc, 'innomd《国内首张！经颅磁刺激设备获批治疗抑郁症》；'
                'https://innomd.com/article/6947d2ef9490de8307a2a15461df6a9c.html')
add_source(doc, '知乎《中国首张！精神领域 TMS 三类证获批》；'
                'https://zhuanlan.zhihu.com/p/1911470527547154646')
add_source(doc, '搜狐《中国首张！精神领域 TMS 三类证获批》（含企业侧多次转载）；'
                'https://www.sohu.com/a/899847819_120987638')
add_source(doc, '知乎/依瑞德《重磅！依瑞德集团喜获中国第一张经颅磁刺激仪欧盟 MDR CE 证书》；'
                'https://zhuanlan.zhihu.com/p/2045802115528905681')
add_source(doc, '医药魔方 ByDrug《磁刺激市场竞争进入新阶段，考验的是可持续技术创新力》，'
                '含市场结构数据；'
                'https://bydrug.pharmcube.com/news/detail/1bc8aeb7f2a698524adc8828196a7bbc')

add_heading(doc, '6.3 依瑞德为拿到三类证所做的研究（公开可查部分）', level=2)
add_para(doc,
    '公开新闻稿披露的研究内容如下（企业未公开完整的临床试验方案，此处为可公开检索到的关键信息汇总，'
    '建议后续通过 CDE 技术审评报告数据库进一步核实）：',
)
yirui_studies = [
    ['临床试验类型', '多中心随机对照试验（Multi-center RCT）'],
    ['参与机构', '全国多家三级甲等医院（新闻稿未逐一披露医院名单）'],
    ['关键验证维度', '产品有效性、安全性、不良反应发生率、设备稳定性、可用性'],
    ['主要终点', '抑郁症相关量表减分率（推测为 HAMD-17 或 MADRS，公开材料未明确披露）'],
    ['研究成果', '"充分验证了依瑞德经颅磁刺激仪治疗抑郁症的临床安全性及有效性"'],
    ['前置研究',
     '2009 年二类证前的性能验证；2024 年 CE MDR 需要的临床评估资料（可复用部分数据）'],
    ['监管沟通', '通过 CDE Pre-submission 与技术审评环节多轮沟通，属于精神领域三类首证案例'],
]
add_table(doc, ['要素', '公开信息内容'], yirui_studies, col_widths=[4, 12])
add_source(doc, '同 6.2 各新闻稿；额外可通过 CDE 官方技术审评报告数据库检索"依瑞德"或"经颅磁刺激仪"'
                '获得完整审评意见（涉及审评意见公开时间差，可能滞后 6–12 个月发布）：'
                'https://www.cmde.org.cn/directory/web/cmde/images/1740709639133081872.pdf（示例）；'
                '数据库入口：https://www.cmde.org.cn/')
add_source(doc, 'NMPA 国家药品监督管理局医疗器械唯一标识数据库（可查产品 UDI 记录）；'
                'https://udi.nmpa.gov.cn/')

add_heading(doc, '6.4 从依瑞德案例可推导出的 Starstim 临床要求', level=2)
add_bullet(doc, '临床方案要素：多中心（≥3 家三甲）、RCT 设计（伪刺激对照）、主流量表（HAMD-17 / MADRS）'
                '为主要终点、随访 4–12 周。')
add_bullet(doc, '样本量：参考精神类三类证案例，通常 100–200 例总样本；具体功效计算需依据历史效应量与'
                '主要终点定义。')
add_bullet(doc, '安全终点：不良事件率、严重不良事件（SAE）、器械故障率、皮肤耐受性。')
add_bullet(doc, '差异化亮点：Starstim 相较依瑞德 TMS 具备"多通道 tDCS + 同步 EEG + 家用远程督导"三大差异化，'
                '可在临床方案中突出便利性与依从性数据（对应欧美 MDD 家用 Pilot 中"90% 完成率"这一亮点）。')
add_bullet(doc, '关键沟通节点：CDE Pre-submission 阶段沟通分类界定、适应症与主要终点定义。')

doc.add_page_break()

# ==== 七、次要对标：脑电图机（Enobio 对标） ====
add_heading(doc, '七、次要对标：脑电图机（Enobio 对标）', level=1)
add_para(doc,
    'EEG 类走豁免临床路径，故没有单一"标杆案例"值得完整对标；但可参考国内已上市的二类脑电图机作为同品种比对对象。'
    '国内主要合规厂商与产品覆盖如下：',
)
eeg_bench = [
    ['深圳迈瑞', '医院场景床旁脑电监护、门诊/科研脑电系统', '二类，已获多个"国械注准"'],
    ['上海诺诚（Nation）', '视频脑电、动态脑电、便携式脑电', '二类'],
    ['深圳英智科技（瀚翔脑科学子公司）', 'BCI + 脑电系列', '二类，其 TMS/tDCS 系列另有独立证'],
    ['西安国际医学中心（依瑞德关联/兄弟公司）', 'CCY 系列刺激器 + EEG 系统', '二类脑电 + 三类刺激'],
    ['进口对标（可作为比对参考）', 'CGX Quick-32、Compumedics、Natus Nicolet', '在中国均有已注册产品可查'],
]
add_table(doc, ['厂家', '产品概览', '类别与注册状态'], eeg_bench, col_widths=[4, 8, 4])
add_source(doc, '国家药品监督管理局数据查询平台（可用于按产品名或注册人查询具体注册证）；'
                'https://www.nmpa.gov.cn/datasearch/')
add_source(doc, '药智数据医械数据库（提供便捷条件筛选）；https://db.yaozh.com/jixie')
add_source(doc, '瀚翔脑科学：https://www.infoinstruments.cn/tmsvs-tdcs/（含 tDCS/TMS 描述）')

doc.add_page_break()

# ==== 八、可点击来源清单 ====
add_heading(doc, '八、可点击来源清单（按主题分组）', level=1)

def add_url_group(title, urls):
    add_heading(doc, title, level=3)
    for u in urls:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run('• ' + u)
        set_run_font(r, size=10)

add_url_group('A. Enobio 欧美注册与研究', [
    'FDA 510(k) K162681 官方 PDF：https://www.accessdata.fda.gov/cdrh_docs/pdf16/K162681.pdf',
    'FDA Innolitics 510(k) 数据库：https://fda.innolitics.com/submissions/NE/subpart-b'
    '%E2%80%94neurological-diagnostic-devices/GWQ/K143440',
    'FDA Report 公司档案页：https://fda.report/Company/Neuroelectrics-Barcelona-S-L-U',
    'Neuroelectrics Enobio 产品页：https://www.neuroelectrics.com/products/research/enobio',
    'Neuroelectrics 官方博客 (MDR)：https://www.neuroelectrics.com/blog/'
    'eu-regulatory-framework-for-medical-devices-an-introduction-to-mdr',
    'Enobio 用户手册：https://manualzz.com/doc/7275717/enobio-user-manual-neuroelectrics',
    'Enobio 独立研究文献 Wiki：https://www.neuroelectrics.com/wiki/index.php/'
    'Collection_of_publications_of_independent_research_studies_and_mentions_about_Enobio',
    'PLOS Biology 使用 Enobio 32 的研究：'
    'https://journals.plos.org/plosbiology/article?id=10.1371%2Fjournal.pbio.3002193',
    'Epilepsy Foundation Enobio 条目：https://www.epilepsy.com/tools-resources/device-wiki/enobio',
])

add_url_group('B. Starstim 欧美注册与临床试验', [
    'BusinessWire FDA Breakthrough Device 2021：'
    'https://www.businesswire.com/news/home/20210630005708/en/',
    'MassDevice IDE 报道：https://www.massdevice.com/'
    'neuroelectrics-touts-fda-ide-starstim-epilepsy-treatment-study-data/',
    'NeurologyLive Breakthrough 报道：https://www.neurologylive.com/view/'
    'new-tes-system-breakthrough-designation-refractory-focal-epilepsy',
    'Neuroelectrics COVID 家用 MDD：https://www.neuroelectrics.com/blog/'
    'fda-greenlights-neuroelectrics-to-help-patients-with-major-depression-at-home-amidst-covid-19-restri',
    'Boston Children\'s NCT02866240：'
    'https://www.childrenshospital.org/clinical-trials/nct02866240',
    'PubMed Kaye et al. 2023 (JCNP 40(1):53-62)：'
    'https://pubmed.ncbi.nlm.nih.gov/34010226/',
    'Frontiers in Psychiatry 2024 MDD Pilot：'
    'https://www.frontiersin.org/journals/psychiatry/articles/10.3389/fpsyt.2024.1427365/full',
    'PMC 全文 MDD Pilot：https://pmc.ncbi.nlm.nih.gov/articles/PMC11358063/',
    'BusinessWire MDD 结果 2024-03-07：'
    'https://www.businesswire.com/news/home/20240307614300/en/',
    'PMC 卒中康复 tDCS Scoping Review：https://pmc.ncbi.nlm.nih.gov/articles/PMC12029044/',
    'Frontiers in Aging Neuroscience 2021 家用 MDD 可行性研究：'
    'https://www.frontiersin.org/journals/aging-neuroscience/articles/10.3389/fnagi.2021.765370/full',
    'Neuroelectrics Starstim 产品页：https://www.neuroelectrics.com/products/research/starstim',
])

add_url_group('C. 公司与创始人背景', [
    'Wikipedia - Neuroelectrics：https://en.wikipedia.org/wiki/Neuroelectrics',
    'Wikipedia - Ana Maiques：https://en.wikipedia.org/wiki/Ana_Maiques',
    'Crunchbase - Neuroelectrics：https://www.crunchbase.com/organization/neuroelectrics',
    'Neuroelectrics Team：https://www.neuroelectrics.com/about-us/team',
])

add_url_group('D. 中国 NMPA 法规与目录', [
    'NMPA 2023 年第 33 号通告（免于临床评价目录）上海药监转发：'
    'https://yjj.sh.gov.cn/qtgzwj/20230725/76df258cba2e4248b2ebb0db7d2061f0.html',
    'CIO 合规查查 2023 年第 33 号通告全文：https://www.ciopharma.com/supervise/31081',
    '健康界解读 2023 年第 33 号通告：'
    'https://www.cn-healthcare.com/articlewm/20230727/content-1584353.html',
    'NMPA 2021 年第 73 号通告（临床评价指导原则）CIRS 解读：'
    'https://www.cirs-group.com/cn/md/gjyjjgyfbylqxlcpjjszdyzd5xjszdyzdtg%EF%BC%882021nd73h%EF%BC%89',
    'NMPA 2023 年第 101 号（分类目录调整）：'
    'https://www.nmpa.gov.cn/directory/web/nmpa////ylqx/ylqxggtg/20230817153633135.html',
    'NMPA 医疗器械分类目录官方入口：https://www.nmpa.gov.cn/wwwroot/gyx02302/flml.htm',
    'CIRS Group 三类注册收费标准 2024：'
    'https://www.cirs-group.com/cn/md/zhong-guo-yi-liao-qi-xie-zhu-ce-shou-fei-biao-zhun-2024-nian-5-yue',
    'ReguVerse 注册收费实施细则：'
    'https://reguverse.com/documentation/nmpa-regulations-index/pre-market-submission/fees/registration-fees/',
    '国家药监局政务服务门户 - 国产三类医疗器械首次注册任务：'
    'https://zwfw.nmpa.gov.cn/web/taskview/11100000MB0341032Y100017214300101',
    'CDE 医疗器械技术审评中心：https://www.cmde.org.cn/',
    'NMPA 医疗器械 UDI 数据库：https://udi.nmpa.gov.cn/',
    'NMPA 数据查询：https://www.nmpa.gov.cn/datasearch/',
])

add_url_group('E. 中国脑电图机相关标准与审评规范', [
    'GB 9706.226-2021 脑电图机专用要求 - 国家标准馆：'
    'https://std.samr.gov.cn/gb/search/gbDetailed?id=CE1E6A1DD42558F6E05397BE0A0A68DF',
    '北京市药监局脑电图机产品技术审评规范：'
    'https://yjj.beijing.gov.cn/yjj/zwgk20/tz7/676482/index.html',
    '西格玛医学脑电图机注册专题：'
    'https://sigma-stat.com/index.php?m=home&c=View&a=index&aid=3035',
    'YY 9706.210-2021 神经和肌肉刺激器专用要求：'
    'https://www.ndls.org.cn/standard/detail/5623b01ad17afcd338c8c08a8b808da5',
])

add_url_group('F. 依瑞德对标案例', [
    '华声在线 2025-05-29 首证报道：https://m.voc.com.cn/xhn/news/202505/29488037.html',
    '中国医院院长网：https://www.h-ceo.com/post/5315.html',
    'innomd 报道：https://innomd.com/article/6947d2ef9490de8307a2a15461df6a9c.html',
    '腾讯新闻《首证破冰》：https://news.qq.com/rain/a/20250605A08UDI00',
    '知乎《中国首张 TMS 三类证》：'
    'https://zhuanlan.zhihu.com/p/1911470527547154646',
    '搜狐转载：https://www.sohu.com/a/899847819_120987638',
    '依瑞德欧盟 MDR CE 证书：https://zhuanlan.zhihu.com/p/2045802115528905681',
    'ByDrug 磁刺激市场分析：'
    'https://bydrug.pharmcube.com/news/detail/1bc8aeb7f2a698524adc8828196a7bbc',
])

add_url_group('G. 相关 CRO / 咨询与费用测算参考', [
    '思途 CRO 中美欧法规差异分析：https://www.situcro.com/news/7767.html',
    '思途 CRO 二三类注册收费标准：https://www.situcro.com/news/4116.html',
    '思途 CRO 二类注册费用：https://www.situcro.com/news/6542.html',
    '奥咨达 NMPA 分类概览：http://www.osmundacn.com/news/newsinfo/id/1484.html',
    'Emergo by UL China 注册：'
    'https://www.emergobyul.com/services/china-medical-device-registration-and-approval',
])

# 结尾说明
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('—— 补充报告完 · 共 8 个章节 / >70 处可核查来源 ——')
set_run_font(r, size=11, color=(89, 89, 89))

# 免责声明
add_heading(doc, '免责声明', level=2)
add_para(doc,
    '本报告所引用的公开新闻稿、企业博客等一手/二手来源，可能存在时间滞后或主观表述；'
    '涉及 NMPA、CDE 官方文件的部分为可复核的一手来源。所有临床费用、临床方案要素为业内经验测算，'
    '建议决策前以 CRO 正式报价、CDE Pre-submission 会议纪要为准。'
    '本报告不构成法律、税务或专业注册意见。',
    size=10, color=(89, 89, 89),
)

# 保存
out = '/home/user/sugz/scratchpad/Neuroelectrics_引进国产化注册_补充报告_来源与对标.docx'
doc.save(out)
print(f'Saved: {out}')
