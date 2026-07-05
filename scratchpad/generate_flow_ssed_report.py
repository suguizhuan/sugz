# -*- coding: utf-8 -*-
"""
FDA PMA P230024 (Flow FL-100 tDCS 治疗抑郁症) SSED 逐项梳理
——用于 Starstim 中国临床方案设计的参照模板
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_run_font(run, name='微软雅黑', size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rpr = r.rPr
    if rpr is None:
        rpr = OxmlElement('w:rPr')
        r.insert(0, rpr)
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def add_heading(doc, text, level=1):
    styles = {
        1: (16, True, (31, 73, 125)),
        2: (13.5, True, (54, 95, 145)),
        3: (12, True, (68, 114, 148)),
    }
    size, bold, color = styles.get(level, (11, True, (0, 0, 0)))
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_para(doc, text, bold=False, indent_cn=True, size=10.5, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent_cn:
        p.paragraph_format.first_line_indent = Pt(21)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(14)
    r1 = p.add_run('【中国注册参考】')
    set_run_font(r1, size=10, bold=True, color=(0, 112, 74))
    r2 = p.add_run(text)
    set_run_font(r2, size=10, color=(89, 89, 89))
    return p


def add_source(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Pt(14)
    r1 = p.add_run('[原文出处] ')
    set_run_font(r1, size=9.5, bold=True, color=(192, 0, 0))
    r2 = p.add_run(text)
    set_run_font(r2, size=9.5, color=(89, 89, 89))
    return p


def add_table(doc, headers, rows, col_widths=None, header_color='1F497D'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        set_cell_bg(hdr[i], header_color)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ''
            p = cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            set_run_font(run, size=9.5)
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


# ============ 生成文档 ============
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ==== 封面 ====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(90)
r = p.add_run('FDA PMA P230024 (Flow FL-100)')
set_run_font(r, size=20, bold=True, color=(31, 73, 125))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('全球首个 tDCS 家用治疗抑郁症')
set_run_font(r, size=18, bold=True, color=(31, 73, 125))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('FDA 三类批准 · 全部研究项目逐项梳理')
set_run_font(r, size=18, bold=True, color=(31, 73, 125))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
r = p.add_run('—— 用于 Neuroelectrics Starstim 中国 NMPA 临床方案设计的参照模板 ——')
set_run_font(r, size=12, color=(89, 89, 89))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
r = p.add_run('原始文件：FDA Summary of Safety and Effectiveness Data (SSED)')
set_run_font(r, size=11, color=(89, 89, 89))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('FDA 批准日期：2025 年 12 月 8 日  |  Breakthrough Device 认定：2022 年 5 月 31 日')
set_run_font(r, size=11, color=(89, 89, 89))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run('梳理时间：2026 年 7 月')
set_run_font(r, size=11, color=(89, 89, 89))

doc.add_page_break()

# ==== 阅读指南 ====
add_heading(doc, '如何使用本参照模板', level=1)
add_para(doc,
    '本文件基于 FDA 于 2025-12-08 批准 Flow FL-100（tDCS 家用抑郁治疗）时公开的《Summary of Safety and '
    'Effectiveness Data，SSED》（PMA 号 P230024），把审批过程中考察的**每一项研究内容**'
    '分类翻译、结构化，并逐项加上"【中国注册参考】"标签，说明该项在 Starstim 中国临床方案中应如何设计。'
)
add_para(doc,
    'Flow FL-100 与 Starstim 高度可比——两者都是 tDCS 家用抑郁治疗设备，都走"FDA Breakthrough → 三类批准"路径，'
    'Flow 于 2019 年在欧盟上市、2025 年获 FDA PMA。这份 SSED 是全球 tDCS 治疗 MDD 的**唯一三类监管路径完整披露**。'
    '中国 NMPA 三类审评虽有本土化差异，但方案要素高度一致，可作为一比一的临床方案骨架直接借鉴。',
)
add_para(doc, '本文件章节对应关系：', bold=True, indent_cn=False)
add_bullet(doc, '第一章 产品与监管概况（对应 SSED §I-VIII）')
add_bullet(doc, '第二章 非临床研究 5 项（对应 SSED §X）')
add_bullet(doc, '第三章 关键临床试验 Empower 详解（对应 SSED §XI，最关键的一章）')
add_bullet(doc, '第四章 补充临床研究 PSYLECT（对应 SSED §XIII，可作为阴性结果的教训案例）')
add_bullet(doc, '第五章 FDA 益处/风险决策逻辑（对应 SSED §XV）')
add_bullet(doc, '第六章 Starstim 中国临床方案建议（本报告的原创综合）')

doc.add_page_break()

# ============ 第一章 产品与监管概况 ============
add_heading(doc, '第一章 产品与监管概况', level=1)
add_source(doc, 'FDA SSED §I–§VIII，PMA P230024 Flow FL-100')

add_heading(doc, '1.1 产品基本信息', level=2)
info_rows = [
    ['通用名', 'Cranial electrotherapy stimulator to treat depression（治疗抑郁症的经颅电刺激器）'],
    ['商品名', 'Flow FL-100'],
    ['配件', 'Flow Clinic Patient Platform（FL-CPP）、Headset Pads（FL-PADS，一次性电极片）'],
    ['产品分类代码', 'JXK（FDA Product Code）'],
    ['申请人', 'Flow Neuroscience AB（瑞典马尔默）'],
    ['PMA 编号', 'P230024'],
    ['FDA 批准日期', '2025 年 12 月 8 日'],
    ['Breakthrough Device 认定', '2022 年 5 月 31 日'],
    ['已上市地域', '欧盟/英国 2019、巴西 2020、澳大利亚 2025'],
    ['是否曾撤市', '否（无安全或有效性相关的撤市记录）'],
]
add_table(doc, ['项目', '内容'], info_rows, col_widths=[4, 12])

add_heading(doc, '1.2 预期用途（Indications for Use）', level=2)
add_para(doc,
    '"Flow FL-100 is intended for the treatment of moderate to severe major depressive disorder (MDD) '
    'in the current episode, either as monotherapy or as an adjunctive treatment, in patients 18 years and '
    'older who are not considered treatment refractory to medication."',
    color=(89, 89, 89),
)
add_para(doc,
    '中文：适用于 18 岁及以上、当前处于中度至重度重度抑郁症（MDD）发作期、'
    '且不属于"药物治疗抵抗性"的患者，可作为单药治疗或辅助治疗使用。',
    bold=True,
)
add_note(doc,
    'Starstim 中国方案建议使用相同或更保守的适应症声称："成人（18–65 岁）中度至重度重度抑郁症发作期的辅助治疗"，'
    '避开"治疗抵抗性抑郁 TRD"这一 CDE 审评难点。'
)

add_heading(doc, '1.3 局限性（Limitations）', level=2)
add_para(doc,
    'FDA 明确列出 Flow FL-100 的评估存在"中等程度的获益不确定性"，主要因素：',
)
add_bullet(doc, '设盲失败（unblinding）：Active 组 77.6% 猜对分组、Sham 组 40.7% 猜对（p=0.012）——设盲不平衡')
add_bullet(doc, '未预设 HDRS-17 分数的"临床显著变化阈值"')
add_bullet(doc, '文献中 tDCS 治疗抑郁的结果存在矛盾')

add_heading(doc, '1.4 禁忌症', level=2)
add_bullet(doc, '电极部位开放性伤口、破损或受损皮肤')
add_bullet(doc, '电极部位有金属颅骨重建物')

add_heading(doc, '1.5 装置描述', level=2)
device_rows = [
    ['刺激参数', '2.0 mA ± 0.05 mA 直流电（tDCS）'],
    ['刺激时长', '每次 30 分钟'],
    ['疗程安排', '前 3 周每周 5 次 + 后 7 周每周 3 次；共 10 周 36 次'],
    ['组件 1 头戴装置', 'Bluetooth（BLE）无线连接；250 mAh 锂聚合物可充电电池（10-15 次/充）；Micro-USB 充电'],
    ['组件 2 电极片', '一次性使用；5 cm 直径圆形；纤维素海绵浸润生理盐水；医用硅胶固定环；铝箔袋成对包装'],
    ['组件 3 手机 App', '引导使用、控制刺激、监测使用记录、发提醒；追踪设备 UDI；每周 MADRS-s 自评'],
    ['组件 4 主服务器', '数据存储 + 处理逻辑后端'],
    ['组件 5 CPP 医生平台', '临床医生 web 端监测患者使用与设置刺激参数'],
    ['使用场景', '家庭（家用）'],
]
add_table(doc, ['要素', '规格'], device_rows, col_widths=[4, 12])
add_note(doc,
    'Starstim 硬件配置更复杂（多通道 tDCS + 同步 EEG），但基础刺激参数（2 mA、30 min）与 Flow 一致，'
    '中国注册技术要求可对标 GB 9706.1 + YY 9706.210 神经肌肉刺激器专用要求。'
)

add_heading(doc, '1.6 替代治疗方案（Alternative Practices）', level=2)
add_para(doc, 'FDA 审评明确列出的抑郁症替代治疗（这些是审评时被拿来做"临床获益 vs 风险"横向对比的对手）：')
add_bullet(doc, '心理治疗：CBT（认知行为治疗）、IPT（人际治疗）、PST（问题解决治疗）')
add_bullet(doc, '药物治疗：SSRIs、SNRIs、非典型抗抑郁药、TCAs、MAOIs、5-HT 调节剂、NMDA 受体拮抗剂')
add_bullet(doc, '器械治疗：TMS（经颅磁）、VNS（迷走神经）、ECT（电休克）、以及基于 App 的数字疗法')
add_note(doc,
    'Starstim 中国注册的临床评价报告（CER）中，需针对上述所有替代疗法做同品种比对分析（尤其是国内首张精神领域 TMS '
    '三类证依瑞德的 CCY 系列），突出 tDCS 的"家用、无手术、低不良事件"三大差异化优势。'
)

add_heading(doc, '1.7 可能的不良反应（Potential AEs）', level=2)
add_para(doc, 'FDA 明确列出的临床研究中观察到的 AE：')
add_bullet(doc, '皮肤干燥、皮肤刺激或红斑（电极部位）')
add_bullet(doc, '刺痛、刺麻、麻木、烧灼感（刺激期间）')
add_bullet(doc, '刺激期间难以集中注意力')
add_bullet(doc, '头痛')
add_bullet(doc, '耳鸣')
add_bullet(doc, '皮肤烧伤')
add_bullet(doc, '嗜睡')
add_bullet(doc, '急性情绪变化 / 抑郁加重')
add_para(doc, '其他文献报告的 tDCS 潜在 AE：')
add_bullet(doc, '心率加速、急性情绪波动、呼吸改变、视物模糊、瘀伤、嗡鸣感、头晕/恶心、疲劳/嗜睡、'
                '潮热、偏头痛、颈痛、皮肤疼痛、生动梦境、狂躁/轻躁狂（其他文献报告）')
add_note(doc,
    'Starstim 方案的"AE 收集清单"可直接采用上述条目作为 tDCS 领域标准 AE list，符合 FDA 审评惯例，'
    '也便于 CDE 审评时的可比性判断。建议使用 tDCS 领域已发表的 AEQ（Adverse Events Questionnaire）作为主 AE 收集工具。'
)

doc.add_page_break()

# ============ 第二章 非临床研究 ============
add_heading(doc, '第二章 非临床研究（5 项）', level=1)
add_source(doc, 'FDA SSED §X（Summary of Non-Clinical Studies）')

add_heading(doc, '2.1 生物相容性（Biocompatibility）', level=2)
bio_rows = [
    ['接触部位', '完整皮肤（intact skin）'],
    ['接触时长', '单次 ≤ 30 分钟；累计使用 ≥ 48 天，累计接触 > 24 小时（分散在数月内）'],
    ['关键材料', 'ABS 塑料、医用硅胶、纤维素海绵（均为医疗器械常用低风险材料）'],
    ['评价依据', 'ISO 10993-1 生物相容性风险管理评价；FDA 的 "Attachment G: Biocompatibility of Certain '
                 'Devices in Contact with Intact Skin" 指南'],
]
add_table(doc, ['项目', '内容'], bio_rows, col_widths=[4, 12])
add_note(doc,
    '中国国标 GB/T 16886 系列（等同 ISO 10993）是必查项。Starstim 走 NMPA 需完成：'
    '（1）细胞毒性 GB/T 16886.5；（2）皮肤致敏 GB/T 16886.10；（3）皮内反应 GB/T 16886.23。'
    '境外已有生物相容性数据可作为引用来源，减少国内测试量。'
)

add_heading(doc, '2.2 电磁兼容与电气安全（EMC & Electrical Safety）', level=2)
add_para(doc, '通过测试的国际标准：')
add_bullet(doc, 'IEC 60601-1 医用电气设备通用安全要求')
add_bullet(doc, 'IEC 60601-1-2 电磁兼容并列标准')
add_bullet(doc, 'IEC 60601-1-11 家用医疗电气设备并列标准（关键！家用场景）')
add_note(doc,
    'Starstim 中国注册对应标准：'
    '（1）GB 9706.1-2020（等同 IEC 60601-1）；'
    '（2）YY 9706.102-2021（等同 IEC 60601-1-2）；'
    '（3）**YY 9706.111（等同 IEC 60601-1-11）家用医疗设备并列标准，如果 Starstim 走家用场景必须做**；'
    '（4）YY 9706.210-2021 神经和肌肉刺激器专用要求。'
    '注册检验必须在国家医疗器械检验所或 CNAS 认可的第三方机构完成。'
)

add_heading(doc, '2.3 软件与网络安全（Software & Cybersecurity）', level=2)
add_para(doc, '关键要点：')
add_bullet(doc, '软件组件：Flow App（患者端）+ CPP 医生门户 + 服务器后端')
add_bullet(doc, '软件安全等级：Major level of concern（因是三类器械）')
add_bullet(doc, '文档要求：完整的软件生命周期文档 + 网络安全文档')
add_note(doc,
    'Starstim 中国方案对应的国内标准：'
    '（1）YY/T 0664-2020 医疗器械软件生存周期过程；'
    '（2）YY/T 1843-2022 医疗器械网络安全；'
    '（3）NMPA《医疗器械软件注册审查指导原则》2022 修订版；'
    '（4）NMPA《医疗器械网络安全注册审查指导原则》2022 修订版。'
    'App 端（患者）和 CPP（医生）都属于 SaMD（Software as Medical Device），需按软件产品单独注册或作为设备一部分注册。'
)

add_heading(doc, '2.4 有效期（Shelf life）', level=2)
add_para(doc, '验证的有效期：3 年（头戴装置 + 电极片）')
add_note(doc,
    'Starstim 的有效期验证通常通过加速老化试验完成（YY/T 0681 系列包装验证 + 加速老化推算）。'
    '预留 6–12 个月做完这项。'
)

add_heading(doc, '2.5 性能测试（Performance Testing）', level=2)
add_para(doc, 'FDA 要求：每一台设备在生产环节都验证输出电流为 2.0 ± 0.05 mA。')
add_note(doc,
    'Starstim 出厂验证需覆盖多通道电流精度、波形完整性、总电荷密度（安全阈值）、皮肤接触阻抗监控等，'
    '这些是 tDCS 器械的关键性能指标。'
)

doc.add_page_break()

# ============ 第三章 关键临床试验 Empower ============
add_heading(doc, '第三章 关键临床试验 Empower（最重要参考）', level=1)
add_source(doc, 'FDA SSED §XI（Summary of Primary Clinical Study）+ 发表于 Nature Medicine 2025;31(1):87-95, '
                'PMID 39433921, DOI 10.1038/s41591-024-03305-y')

add_heading(doc, '3.1 试验总览', level=2)
empower_rows = [
    ['试验名称', 'Empower Study'],
    ['正式发表', 'Woodham et al. Home-based tDCS treatment for MDD: a fully remote phase 2 RCT. '
                'Nature Medicine 2025;31(1):87-95'],
    ['IDE 编号', 'G210328'],
    ['试验持续期', '2022 年 6 月 – 2023 年 8 月（患者治疗期）'],
    ['数据库锁定', '2023 年 8 月'],
    ['样本量', 'ITT = 174；mITT = 173（1 例随机后未使用设备）'],
    ['地域', '英国（英格兰/威尔士）115 例 + 美国（德州）59 例'],
    ['中心数', '2 个研究中心（伦敦国王学院附属 King\'s College 医院 + 德州大学 UTHealth 休斯顿分校）'],
    ['随机比例', '1:1（active vs sham）'],
    ['设盲', '双盲（患者 + 评估者）'],
    ['执行模式', '全远程家用（fully remote at-home）'],
    ['总时长', '20 周（10 周盲态 + 10 周开放标签）'],
    ['监督方式', '首次现场指导；此后通过 CPP 医生端远程监督'],
]
add_table(doc, ['项目', '内容'], empower_rows, col_widths=[4, 12])
add_note(doc,
    '**关键启示**：Flow 只用了 2 个中心 174 例患者就拿下 FDA 三类证。中国 CDE 惯例是 3–5 中心、150–300 例，'
    '因此 Starstim 中国临床可以按"3 中心 150 例（100 active + 50 sham 或 1:1）"设计——这是"够用且经济"的样本量。'
    'Flow 的"全远程家用"设计极大降低了 CRC 与随访成本，是 Starstim 应重点学习的模式。'
)

add_heading(doc, '3.2 治疗方案', level=2)
tx_rows = [
    ['盲态期（Week 0-10）', 'Active：2.0 mA tDCS × 30 min，共 36 次'],
    ['- 前 3 周', '每周 5 次'],
    ['- 后 7 周', '每周 3 次'],
    ['Sham（对照）方案', '30 秒内电流缓升至 1 mA，15 秒内缓降至 0 mA；30 分钟结束前 45 秒重复一次'],
    ['开放标签期（Week 10-20）', 'Sham 组转 Active 治疗 10 周；原 Active 组维持治疗每周 3 次'],
    ['首次刺激监督', '由研究人员现场指导'],
    ['后续监督', 'CPP 医生端远程监督（研究员定期查看使用记录）'],
]
add_table(doc, ['阶段/参数', '内容'], tx_rows, col_widths=[4, 12])
add_note(doc,
    'Starstim 中国方案可直接采用相同或简化的治疗强度（2 mA × 30 min × 36 次），但需选择"多中心均可支持的物流方案"。'
    '中国国情下"全远程家用"仍有争议（CDE 通常希望有中心访视），建议改良为"前 4 次在院指导 + 后续家用远程督导 + Week 5/10/20 现场访视"。'
)

add_heading(doc, '3.3 入选标准（15 条，Inclusion Criteria）', level=2)
incl = [
    '年龄 ≥ 18 岁',
    'DSM-V 诊断为单相 MDD 当前发作期',
    'HDRS-17 评分 ≥ 16 分',
    '入组前 6 周内：未服抗抑郁药 或 使用稳定的抗抑郁药方案（同意维持不变）',
    '如接受心理治疗，入组前 6 周维持稳定',
    '拥有稳定网络连接',
    '拥有 Android 5.0+ / iOS 12+ 智能手机且能熟练使用',
    '当前居住地在英格兰/威尔士（英国）或德州（美国）',
    '当前由精神科医生或初级保健医师管理，同意定期评估',
    '同意研究人员与其医疗提供者沟通',
    '提供至少 2 位居住在 60 分钟车程内的紧急联系人',
    '签署知情同意书',
    '愿意配合所有试验流程',
    '在试验期间维持所有入选标准',
    '同意参与"自杀风险管理协议"',
]
for i, item in enumerate(incl, 1):
    add_numbered(doc, item)
add_note(doc,
    'Starstim 中国方案的入选标准可简化到 10–12 条，保留 DSM-5 诊断、HDRS-17 ≥ 16、'
    '药物治疗稳定/未用、可智能手机操作、稳定联系方式几项核心即可。'
    '"60 分钟车程紧急联系人"这一条是 FDA 特有的家用安全要求，中国也可保留一条简化版。'
)

add_heading(doc, '3.4 排除标准（35 条，Exclusion Criteria）', level=2)
add_para(doc, 'FDA 审评的排除标准非常严格，共 35 条。以下按类别归类：')

excl_groups = [
    ('精神/情感状态', [
        '当前狂躁状态（YMRS）或精神病状态（MINI）',
        '其他 MDD 干预治疗（除稳定抗抑郁药与心理治疗）',
        '治疗抵抗性抑郁 TRD（≥2 次足量足疗程抗抑郁药无效）',
        'C-SSRS Triage 中问题 4/5/6 回答"是"',
        '有自杀行为住院史',
        '曾诊断 OCD、双相 1 或 2 型',
        '当前诊断为主要焦虑障碍、PTSD、恐惧症、厌食/贪食、惊恐障碍、人格障碍（活动症状）',
        '有狂躁或精神病病史',
        '曾用氯胺酮/艾司氯胺酮治疗抑郁',
        '曾因抑郁住院',
        '曾接受抑郁精神外科手术',
    ]),
    ('神经系统', [
        '既往有 ECT、TMS、CES、tDCS、DBS 或其他脑刺激史',
        '结构性脑病变（卒中、皮质下病变异常等）',
        '脑内 / 颅骨植入物（DBS 等）',
        '颅骨缺损',
        '癫痫或痉挛发作史（含戒断/诱发）',
        '头部有金属残片或铁磁材料',
        '曾诊断自闭症谱系障碍',
        '认知障碍（含痴呆）',
        '帕金森病或其他运动障碍',
        '有慢性偏头痛',
    ]),
    ('医学与生活方式', [
        '维生素或激素缺乏可能类似情绪障碍',
        '慢性重度失眠（< 4 h/晚）或抑郁继发于慢性失眠/睡眠呼吸暂停',
        '任何可能影响完成问卷的疾病',
        '当前物质滥用（入组前 <1 周）',
        '当前使用影响皮质兴奋性的药物（如苯二氮䓬、抗癫痫药）',
        '当前有酒精或苯二氮䓬戒断症状',
        '心肌梗死、CABG、CHF 或其他心脏病史',
        '慢性烟民（终生 > 100 支且过去 7 天每天吸烟）',
        '孕期、哺乳期、试图怀孕/哺乳',
    ]),
    ('研究执行相关', [
        '当前为囚犯',
        '正参与或过去 90 天参与过其他临床试验（可能干扰结果）',
        '任何可能影响随访完成或 App 自评的医学/环境情况',
        '研究者判断影响装置安全/性能评估的任何情况',
        '在试验期间违反排除标准者将退出研究',
    ]),
]
for group_name, items in excl_groups:
    add_para(doc, f'{group_name}（{len(items)} 条）', bold=True, indent_cn=False)
    for it in items:
        add_bullet(doc, it, level=1)
add_note(doc,
    '**这份排除标准是全球 tDCS 治疗 MDD 三类审批级别的黄金标准**。'
    'Starstim 中国方案可直接复制 90% 以上条目；针对中国特色补充：'
    '（1）明确排除 TRD（依瑞德 CCY-IV 三类证也采取相同策略）；'
    '（2）明确"神经调控史阳性排除"以避免污染；'
    '（3）建议再加"HAMD 之外配合 CGI-S 与 CGI-I 评估"以符合国内精神科惯例。'
)

add_heading(doc, '3.5 访视计划（Schedule of Procedures）', level=2)
visit_rows = [
    ['筛选/基线（Wk -3 to -1）', '知情同意、筛选视频访谈、MDD 诊断评估、HDRS-17、MADRS、C-SSRS、YMRS、'
                                'HAM-A、EQ-5D-3L、RAVLT、TAQ、既往药物记录'],
    ['Week 0（随机）', '随机化、技术上机指导、初诊访视、MADRS-s（App 自评）、AE 记录开始'],
    ['Week 1（±3d）', 'HDRS-17、MADRS、MADRS-s、C-SSRS、YMRS、AE、合并用药'],
    ['Week 4（±3d）', 'HDRS-17、MADRS、MADRS-s、C-SSRS、YMRS、AE、合并用药'],
    ['Week 7（±3d）', 'HDRS-17、MADRS、MADRS-s、C-SSRS、YMRS、AE、合并用药'],
    ['**Week 10（±3d，主要终点访视）**', 'HDRS-17、MADRS、MADRS-s、C-SSRS、YMRS、HAM-A、EQ-5D-3L、RAVLT、'
                                        'TAQ、AEQ、Healthcare Visit Survey、盲态揭盲'],
    ['Week 10 开放标签期开始', '揭盲：Sham 组开始 Active；Active 组继续 3 次/周维持治疗'],
    ['Week 20（±3d，试验结束）', '重复主要终点访视的全部量表'],
    ['提前终止（Weeks 1-20）', '重复主要访视全部量表 + 中止访谈（可选）'],
]
add_table(doc, ['访视时点', '关键流程/量表'], visit_rows, col_widths=[5, 11])
add_note(doc,
    'Starstim 中国方案可采用同样的访视密度（6 次评估：基线/Wk1/4/7/10/20），符合 CDE 对精神科三类临床的标准要求。'
    '建议加入 CGI-S/CGI-I 量表（中国精神科惯用）。'
)

add_heading(doc, '3.6 评价指标（Endpoints）', level=2)
add_heading(doc, '主要终点', level=3)
add_para(doc,
    '**Week 10 时，Active vs Sham 组间 HDRS-17 均分改变量的调整后差值（Adjusted mean group difference）**',
    color=(192, 0, 0),
)
add_para(doc, '统计假设：H0：d_flow - d_sham ≤ 0；Ha：d_flow - d_sham > 0（单侧检验，α = 0.025）')

add_heading(doc, '次要终点（9 项）', level=3)
sec_rows = [
    ['1a', 'Week 10 HDRS-17 组间应答率差异', '入 Hochberg 多重性控制'],
    ['1b', 'Week 10 HDRS-17 组间缓解率差异', '入 Hochberg 多重性控制'],
    ['2a', 'Week 10 MADRS 均分改变差异', ''],
    ['2b', 'Week 10 MADRS 组间应答率差异', ''],
    ['2c', 'Week 10 MADRS 组间缓解率差异', ''],
    ['3a', 'Week 10 MADRS-s 均分改变差异', ''],
    ['3b', 'Week 10 MADRS-s 组间应答率差异', ''],
    ['3c', 'Week 10 MADRS-s 组间缓解率差异', ''],
    ['4', 'Week 10 EQ-5D-3L 生活质量改善差异', '入 Hochberg 多重性控制'],
]
add_table(doc, ['编号', '内容', '备注'], sec_rows, col_widths=[1.5, 10.5, 4])

add_heading(doc, '关键定义', level=3)
add_bullet(doc, '应答（Response）：从基线到 Week 10 症状评分下降 ≥ 50%')
add_bullet(doc, '缓解（Remission）：Week 10 时 HDRS-17 ≤ 7；或 MADRS ≤ 10；或 MADRS-s ≤ 12')
add_note(doc,
    'Starstim 中国方案的主要终点建议直接采用"Week 8 或 Week 10 HAMD-17 减分率"（或组间均分差异），'
    '这是国内精神科临床的标准。次要终点应保留 MADRS、CGI-I、生活质量量表（可改用中国常用的 SF-36 或 EQ-5D-5L）。'
)

add_heading(doc, '3.7 统计方法', level=2)
add_bullet(doc, '分析人群：mITT（modified ITT）= 至少接受 1 次刺激治疗的受试者（n=173）')
add_bullet(doc, '主要分析模型：Mixed Model for Repeated Measures（重复测量混合模型）')
add_bullet(doc, '缺失数据处理：Fully Conditional Specification (FCS) 多重插补，20 个插补数据集，Rubin\'s Rule 合并')
add_bullet(doc, '协变量：年龄、性别、基线是否心理治疗、基线是否用抗抑郁药、治疗组')
add_bullet(doc, '多重性控制：Hochberg 方法（对 3 个预设次要终点：1a、1b、4）')
add_bullet(doc, '成功判定：单侧 p < 0.025')
add_note(doc,
    'Starstim 中国方案的统计部分可直接采用相同方法。CDE 尤其看重"预设的分析计划（SAP）"，'
    '建议在方案启动前完成 SAP 并锁定，且盲态审核（Blind Data Review）后再解盲。'
    '缺失数据插补方案应写在 SAP 中，避免临时决定。'
)

add_heading(doc, '3.8 关键疗效结果（Empower 报告值）', level=2)
result_rows = [
    ['HDRS-17 均分改变', '-9.4', '-7.1', '-2.3', 'Cohen\'s d = 0.37', '0.012 ★'],
    ['HDRS-17 应答率', '54.4%', '26.9%', '27.5%', 'OR = 3.25', '0.001 ★'],
    ['HDRS-17 缓解率', '44.9%', '21.8%', '23.1%', 'OR = 2.93', '0.004 ★'],
    ['MADRS 均分改变', '-11.3', '-7.7', '-3.6', 'Cohen\'s d = 0.41', '0.006 ★'],
    ['MADRS 应答率', '63.0%', '31.6%', '31.4%', 'OR = 3.70', '<0.001 ★'],
    ['MADRS 缓解率', '57.5%', '29.4%', '28.1%', 'OR = 3.26', '0.002 ★'],
    ['MADRS-s 均分改变', '-9.9', '-6.2', '-3.7', 'Cohen\'s d = 0.41', '0.009 ★'],
    ['MADRS-s 应答率', '49.1%', '24.0%', '25.1%', 'OR = 3.06', '0.004 ★'],
    ['MADRS-s 缓解率', '53.8%', '23.4%', '30.4%', 'OR = 3.83', '0.002 ★'],
    ['EQ-5D-3L 均分改变', '0.08', '0.06', '0.02', '—', '0.326 ✗'],
]
add_table(doc, ['终点（Week 10）', 'Active', 'Sham', 'Δ', '效应量', 'p 值'], result_rows,
          col_widths=[4, 2, 2, 2, 3, 3])
add_para(doc,
    '★ = 达到主要或预设次要终点显著性；✗ = 未达显著性。'
    '**主要终点（HDRS-17 均分改变）达到显著性；EQ-5D-3L 未达显著性但 FDA 认为该量表本身对轻度抑郁不敏感。**',
    color=(89, 89, 89), indent_cn=False,
)
add_note(doc,
    'Starstim 中国方案的样本量估算，可以直接以 Flow 的 Cohen\'s d = 0.37 作为参照效应量。'
    '按 α=0.025 单侧、power=80%、两组 1:1、d=0.37 计算，每组约需 116 例，总样本 232 例。'
    '若采用 CDE 常见的 α=0.05 双侧标准，样本可减到约 150 例总样本（可行）。'
)

add_heading(doc, '3.9 亚组分析', level=2)
subgroup_rows = [
    ['单药治疗（Monotherapy）', 'n=52（23 vs 29）', '-9.7 vs -5.9（Δ=-3.8，d=0.56，p=0.0027）',
     '47.8% vs 13.8%（OR=5.76，p<0.0001）', '34.8% vs 13.8%（OR=3.29，p=0.0027）'],
    ['辅助治疗（合用抗抑郁药）', 'n=109（56 vs 53）', '-10.6 vs -7.8（Δ=-2.8，d=0.46，p=0.019）',
     '63.2% vs 35.4%（OR=3.14，p=0.010）', '53.1% vs 25.5%（OR=3.33，p=0.009）'],
]
add_table(doc, ['亚组', 'n', 'HDRS-17 均分改变', '应答率', '缓解率'], subgroup_rows,
          col_widths=[3.5, 2, 3.5, 3.5, 3.5])
add_para(doc,
    '**关键结论**：无论单药还是辅助治疗，Active 组均显著优于 Sham。'
    '性别、年龄、种族、中心、基线 HDRS-17 均未见交互作用，数据可合并分析。',
    bold=True,
)

add_heading(doc, '3.10 设盲评估（重要 FDA 关切）', level=2)
blind_rows = [
    ['真实分组 Active', '77.6% 猜对（Active）', '22.4% 猜错（Sham）'],
    ['真实分组 Sham', '40.7% 猜对（Sham）', '59.3% 猜错（Active）'],
]
add_table(doc, ['真实组', '猜"Active"', '猜"Sham"'], blind_rows, col_widths=[5, 5.5, 5.5])
add_para(doc,
    '组间差异 p=0.012——**设盲失败**（active 组显著感知到自己在接受真实治疗）。'
    'FDA 通过"correct guess rate 调整分析"（CGR）进一步验证：即使假设猜测完全平衡，'
    'HDRS-17 组间差值仍为 -2.0 至 -2.12 分，与原始 -2.3 分接近，因此结论不变。',
)
add_note(doc,
    '**这是 tDCS 类临床方案设计的最大挑战**。中国方案可考虑：'
    '（1）使用相同的标准 sham 方案（升至 1 mA 再降至 0），并在方案中明确"设盲评估"作为独立次要终点；'
    '（2）预先规划 CGR 敏感性分析并写入 SAP；'
    '（3）注意受试者是否可感知刺激强度差异，纳入"感知强度问卷"作为盲态质控指标。'
    '这些前置准备可在 CDE 审评时主动解释，避免"审评发补"。'
)

add_heading(doc, '3.11 安全性主要发现', level=2)
add_bullet(doc, '**无严重不良事件（SAE）归因于设备**——173 例仅 1 例高血压 SAE（active 组），审评认为与设备无关')
add_bullet(doc, '不良事件发生率：Active 组 70.1% vs Sham 组 59.3%')
add_bullet(doc, '最常见 AE（Active% vs Sham%）：皮肤红斑 63.5% vs 18.5%；瘙痒 50.6% vs 43.2%；'
                '烧灼感 43.5% vs 38.3%；头痛 42.4% vs 35.8%；头皮痛 21.2% vs 12.3%')
add_bullet(doc, '2 例一级皮肤烧伤（Active 组），均因电极片干燥导致，休息数日后恢复')
add_bullet(doc, '无狂躁、轻躁狂、自杀相关事件')
add_note(doc,
    'Starstim 中国方案的 AE 收集清单可完全参照 Flow 的 Adverse Events Questionnaire（AEQ），'
    '按 MedDRA 编码上报。皮肤类 AE 是 tDCS 类必然出现的可预期风险，可在知情同意书中充分告知。'
)

add_heading(doc, '3.12 开放标签期结果', level=2)
add_bullet(doc, 'Active 组 55/87 完成开放标签期（63.2%）；Sham 组 56/86 完成')
add_bullet(doc, 'Active 组 Week 10 应答者 31 例中 30 例（97%）在 Week 20 仍为应答者——**疗效可维持**')
add_bullet(doc, 'Active 组 Week 10 非应答者 24 例中 9 例（38%）在 Week 20 转为应答者——**延长治疗可能改善**')
add_bullet(doc, '开放标签期无 SAE')

doc.add_page_break()

# ============ 第四章 补充临床研究 ============
add_heading(doc, '第四章 补充临床研究：PSYLECT（阴性结果的教训）', level=1)
add_source(doc, 'FDA SSED §XIII；Borrione L, et al. Home-Use tDCS for MDD: A Randomized Clinical Trial. '
                'JAMA Psychiatry 2024;81(4):329-337, DOI 10.1001/jamapsychiatry.2023.4948；NCT04889976')

add_heading(doc, '4.1 试验概况', level=2)
psy_rows = [
    ['试验名称', 'PSYLECT（Portable Transcranial Electrical Stimulation and Internet-Based Behavioral Therapy for MDD Study）'],
    ['NCT 编号', 'NCT04889976'],
    ['中心', '巴西圣保罗大学附属医院（Hospital Universitário，单中心）'],
    ['设计', '单中心 3 臂 RCT 双盲 sham 对照'],
    ['样本量', '210 例入组（180 女、平均 38.9 岁），199 例完成'],
    ['3 臂设置', '（1）双阳性 tDCS + iBT；（2）tDCS + sham iBT；（3）双 sham'],
    ['刺激方案', '2 mA × 30 min × 15 个连续工作日 + 之后 3 周每周 2 次；Sham 为 1 mA × 90 s'],
    ['主要终点', 'Week 6 HDRS-17 分数改变'],
]
add_table(doc, ['项目', '内容'], psy_rows, col_widths=[4, 12])

add_heading(doc, '4.2 结果（阴性）', level=2)
add_bullet(doc, '双阳性 vs 仅 tDCS：Cohen\'s d = 0.05（95%CI −0.48 – 0.58，p=0.86）')
add_bullet(doc, '双阳性 vs 双 sham：Cohen\'s d = -0.20（95%CI −0.73 – 0.34，p=0.47）')
add_bullet(doc, '仅 tDCS vs 双 sham：Cohen\'s d = -0.25（95%CI −0.76 – 0.27，p=0.35）')
add_bullet(doc, '皮肤红斑、烧灼感在 tDCS 组更常见')
add_bullet(doc, '1 例非致死自杀企图（仅 tDCS 组）')
add_bullet(doc, '**结论：无监督的家用 tDCS + 数字心理干预未证明优于 sham**')

add_heading(doc, '4.3 PSYLECT 阴性结果的教训', level=2)
add_note(doc,
    'FDA 审评中把 PSYLECT 列为"文献冲突结果"之一，是获益不确定性的来源之一。'
    'Starstim 中国方案应避开 PSYLECT 的失败点：'
    '（1）不要"无监督家用"——保留 CPP 医生端远程监督；'
    '（2）不要合并数字心理干预作为主要变量——单独评估 tDCS 效果；'
    '（3）不要单中心——多中心至少 3 家；'
    '（4）疗程要足够——PSYLECT 只 6 周略短，建议按 Empower 的 10 周设计。'
)

doc.add_page_break()

# ============ 第五章 FDA 益处/风险决策逻辑 ============
add_heading(doc, '第五章 FDA 益处/风险决策逻辑', level=1)
add_source(doc, 'FDA SSED §XV（Conclusions Drawn from Preclinical and Clinical Studies）')

add_heading(doc, '5.1 有效性结论', level=2)
add_bullet(doc, 'HDRS-17 组间差 2.3 分虽达统计显著（p=0.012），但医学文献缺乏"临床显著变化阈值"共识，因此 FDA 只认为"支持有效性可能性"（probable benefit）而非"确证有效"')
add_bullet(doc, '预设的 HDRS-17 应答/缓解次要终点（多重性调整后）均显著')
add_bullet(doc, 'MADRS 和 MADRS-s 次要终点在应答率和缓解率上均显著支持')
add_bullet(doc, '单药与辅助治疗两亚组均显著（monotherapy Cohen\'s d = 0.56，adjunctive d = 0.46）')

add_heading(doc, '5.2 安全性结论', level=2)
add_bullet(doc, 'FL-100 呈现"可能的风险"，但风险等级低')
add_bullet(doc, '风险总体轻度、短暂，主要为皮肤问题与短暂头痛')
add_bullet(doc, '电极片重复使用或干燥可致皮肤烧伤——已纳入警示')

add_heading(doc, '5.3 益处-风险决定', level=2)
add_para(doc,
    'FDA 的决定："The totality of the evidence demonstrates FL-100 provides probable benefit... '
    'sufficient to outweigh its probable risk."（综合证据显示，FL-100 提供的可能获益足以超过其可能风险。）',
)
add_para(doc, '**FDA 特别考虑的加分项**：', bold=True, indent_cn=False)
add_bullet(doc, '临床依从性高（<15% 流失率）')
add_bullet(doc, 'RCT 覆盖 2 个不同地域，支持结果外推')
add_bullet(doc, '单药和辅助治疗均显示可能获益')
add_bullet(doc, '家用使用促进治疗可及性，为药物治疗低耐受患者提供替代')
add_bullet(doc, '欧洲市场长期使用积累的上市后监测数据支持风险特征已被充分表征')
add_note(doc,
    '**决策启示**：Starstim 中国注册时，重点向 CDE 展示以下要素：'
    '（1）Neuroelectrics 已积累的 CE 上市后监测数据（安全性长期证据）；'
    '（2）多中心均一性证据；'
    '（3）单药与辅助两个亚组的一致性；'
    '（4）家用便利性和依从性数据；'
    '（5）与 TMS/ECT 相比的"低风险、无麻醉、无手术"优势。'
)

doc.add_page_break()

# ============ 第六章 Starstim 中国方案建议 ============
add_heading(doc, '第六章 Starstim 中国临床方案建议（一比一映射）', level=1)

add_heading(doc, '6.1 推荐方案框架（一句话总结）', level=2)
add_para(doc,
    '**采用"Empower + 中国本土化"混合方案：3 中心多中心 RCT、150 例总样本、10 周盲态 + 10 周开放标签、'
    '主要终点 HAMD-17 均分变化、次要终点采用 MADRS + CGI-I，允许单药或辅助（合并抗抑郁药）两种用药背景。**',
    bold=True, color=(0, 112, 74),
)

add_heading(doc, '6.2 方案要素一览', level=2)
plan_rows = [
    ['适应症', '18–65 岁 MDD 当前发作期（HAMD-17 ≥ 16），排除 TRD'],
    ['中心数', '3 家三甲医院精神科（推荐北京安定 / 上海精卫 / 广州脑科 或同级中心）'],
    ['样本量', '150 例总（1:1 随机）'],
    ['统计功效', 'α=0.025 单侧，power=80%，Cohen\'s d=0.4（基于 Empower 0.37 保守估计）'],
    ['随机比例', '1:1 Active vs Sham'],
    ['盲态期', '10 周（36 次治疗：前 3 周每周 5 次 + 后 7 周每周 3 次）'],
    ['开放标签期', '10 周（Sham 转 Active；Active 维持 3 次/周）'],
    ['刺激参数', '2.0 mA tDCS × 30 min/次；Sham 采用标准 ramp-up/down 方案'],
    ['执行模式', '"半远程"：Wk 0/4/10/20 现场访视 + 家用为主'],
    ['主要终点', 'Week 10 HAMD-17 均分改变（Active vs Sham 组间差异）'],
    ['次要终点', 'HAMD-17 应答率/缓解率、MADRS、CGI-I、CGI-S、EQ-5D-5L、安全性 AEQ'],
    ['统计方法', 'MMRM 混合模型 + FCS 多重插补 + Hochberg 多重性控制'],
    ['安全监测', 'DSMB 数据安全监察委员会；每周 C-SSRS 自杀评估；AEQ AE 收集'],
    ['预计时长', '入组 12 个月 + 治疗随访 10 个月 + 统计报告 3 个月 = 25 个月'],
    ['预计总成本', '450–600 万元（不含设备与产品技术要求）'],
]
add_table(doc, ['要素', '推荐内容'], plan_rows, col_widths=[4, 12])

add_heading(doc, '6.3 关键风险点与预案', level=2)
risk_rows = [
    ['设盲失败', 'Empower 显示 Active 组 77.6% 猜对分组', '预设 CGR 敏感性分析写入 SAP；受试者"感知强度问卷"作为质控'],
    ['样本流失', 'Empower 流失 14.3%（合理）', '中国方案预留 20% 流失，实际入组 180 例；家用便利性降低流失'],
    ['CDE 审评发补', 'CDE 可能对设盲、多重性、缺失数据处理提问', 'Pre-submission 前置沟通；预留 6 个月发补窗口'],
    ['单中心变异', '3 个中心可能存在异质性', '预设中心作为分层变量；中心-治疗交互检验'],
    ['与依瑞德 TMS 竞争', '国内 TMS 三类证已开始下沉', '突出 tDCS 家用便利、无麻醉、依从性高的独有优势'],
]
add_table(doc, ['风险', '来源', '缓释措施'], risk_rows, col_widths=[3, 5, 8])

add_heading(doc, '6.4 关键参考文献（供 CER 引用）', level=2)
refs = [
    'Woodham RD, et al. Home-based tDCS treatment for major depressive disorder: a fully remote phase 2 '
    'randomized sham-controlled trial. Nat Med. 2025;31(1):87-95. PMID: 39433921. DOI: 10.1038/s41591-024-03305-y.'
    ' — Empower 试验主论文，Starstim 中国方案的一比一参考',

    'Borrione L, et al. Home-Use Transcranial Direct Current Stimulation for the Treatment of a Major '
    'Depressive Episode: A Randomized Clinical Trial. JAMA Psychiatry. 2024;81(4):329-337. '
    'DOI: 10.1001/jamapsychiatry.2023.4948. — PSYLECT 阴性教训',

    'Brunoni AR, et al. Trial of Electrical Direct-Current Therapy versus Escitalopram for Depression. '
    'N Engl J Med. 2017;376(26):2523-2533. PMID: 28657871. — ELECT-TDCS 经典对头对比',

    'Matsumoto H, Ugawa Y. Adverse events of tDCS and tACS: A review. Clin Neurophysiol Pract. 2016;2:19-25. '
    'PMID: 30214966. — tDCS 安全性系统综述',

    'Brunoni AR, et al. A systematic review on reporting and assessment of adverse effects associated with '
    'transcranial direct current stimulation. Int J Neuropsychopharmacol. 2011;14(8):1133-45. PMID: 21320389. '
    '— AEQ 量表来源',

    'Szigeti B, et al. The difference between "placebo group" and "placebo control": a case study in '
    'psychedelic microdosing. Sci Rep. 2023;13:12107. DOI: 10.1038/s41598-023-34938-7. — CGR 设盲调整方法',
]
for i, ref in enumerate(refs, 1):
    add_numbered(doc, ref)

doc.add_page_break()

# ============ 附录 ============
add_heading(doc, '附录：Empower 试验完整量表清单', level=1)
scale_rows = [
    ['HDRS-17', '17 项汉密尔顿抑郁量表', '主要终点', '临床评估'],
    ['MADRS', '蒙哥马利-阿斯伯格抑郁评定量表', '次要终点', '临床评估'],
    ['MADRS-s', 'MADRS 患者自评版', '次要终点', 'App 自评'],
    ['C-SSRS', '哥伦比亚自杀严重程度评定量表', '安全监测', '每次访视'],
    ['YMRS', 'Young 躁狂评定量表', '安全监测', '每次访视'],
    ['HAM-A', 'Hamilton 焦虑量表', '基线 + Wk 10/20', '临床评估'],
    ['EQ-5D-3L', 'EuroQol 5 维 3 级生活质量量表', '次要终点', '基线 + Wk 10/20'],
    ['RAVLT', 'Rey 听觉词语学习测试', '认知功能', '基线 + Wk 10/20'],
    ['SDMT', '符号数字模式测试', '认知功能（选做）', ''],
    ['TAQ', '治疗接受性问卷', '受试者体验', '基线 + Wk 10/20'],
    ['AEQ', '不良事件问卷', '安全终点', 'Wk 10 + 20'],
    ['MINI', 'MINI 国际神经精神访谈', '基线诊断', '筛选期'],
]
add_table(doc, ['缩写', '中文名称', '用途', '访视时点'], scale_rows,
          col_widths=[2.5, 6, 3.5, 4])

add_para(doc, '')
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('—— 报告完 · 全部内容基于 FDA PMA P230024 SSED（35 页）逐项梳理 ——')
set_run_font(r, size=11, color=(89, 89, 89))

# 免责
add_heading(doc, '免责声明', level=2)
add_para(doc,
    '本报告内容全部来源于 FDA 于 2025-12-08 公开发布的 PMA P230024 Summary of Safety and Effectiveness Data '
    '（SSED）原文；中文翻译由 AI 完成，专业术语建议由 Starstim 项目组的精神科顾问二次审核。'
    '"【中国注册参考】"标签内容为本报告作者基于中国 NMPA 现行法规与临床惯例的经验建议，不构成法律或注册意见。'
    '正式方案启动前请与 CDE 通过 Pre-submission 会议确认关键要素。',
    size=10, color=(89, 89, 89),
)

# 保存
out = '/home/user/sugz/scratchpad/FDA_P230024_Flow_FL100_临床方案参考.docx'
doc.save(out)
print(f'Saved: {out}')
